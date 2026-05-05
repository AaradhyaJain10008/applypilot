import os
import json
import math
import re
import csv
import time
import threading
import uuid
import ssl
from datetime import datetime, timedelta, time as dt_time
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from pathlib import Path
from urllib import error as url_error
from urllib import request as url_request
from urllib.parse import quote
import io
import certifi
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import google.generativeai as genai
from google.api_core import retry as gapi_retry
from send_email import send_email
from log_job import log_job
from config_loader import load_app_settings, load_profile, load_resume_personas
from services.apify_client import ApifyClientService, ApifySafeStop, ApifyConfigError
from services.resume_text import extract_pdf_text, ResumePdfError

# override=True ensures that edits to .env take effect on every Flask reload,
# rather than being shadowed by stale values inherited from the parent process env.
load_dotenv(override=True)

app = Flask(__name__)
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))

# Profile / persona / settings load from config/*.json (with .example.json fallback).
PROFILE = load_profile()
PERSONAS = load_resume_personas()
APP_SETTINGS = load_app_settings()

# Feature flag: set to False to disable/remove follow-up chat quickly.
# Now driven by config/app_settings.json -> features.enable_followup_agent.
ENABLE_FOLLOWUP_AGENT = APP_SETTINGS.feature("enable_followup_agent", True)
ENABLE_EMAIL_SENDING = APP_SETTINGS.feature("enable_email_sending", True)
ENABLE_EMAIL_SCHEDULING = APP_SETTINGS.feature("enable_email_scheduling", True)

# Configure Gemini
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    # Force REST transport: gRPC channels do not survive Flask's debug reloader fork
    # and start throwing "503 failed to connect / FAILED_PRECONDITION / Operation not permitted"
    # against local loopback addresses. REST is fork-safe and matches plain HTTPS behavior.
    genai.configure(api_key=api_key, transport="rest")
AI_PROVIDER = os.getenv("AI_PROVIDER", "groq").strip().lower()

# --- Gemini (Google) -----------------------------------------------------
GENAI_MODEL = os.getenv("GEMINI_MODEL", "models/gemini-2.0-flash")
GEMINI_MODEL_CANDIDATES = [
    m.strip() for m in os.getenv("GEMINI_MODEL_CANDIDATES", "").split(",") if m.strip()
]

# --- Groq (free tier, OpenAI-compatible, blazing fast) -------------------
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1").rstrip("/")
GROQ_MODEL_ANALYZE = os.getenv("GROQ_MODEL_ANALYZE", "llama-3.3-70b-versatile")
GROQ_MODEL_DEFAULT = os.getenv("GROQ_MODEL_DEFAULT", "llama-3.3-70b-versatile")
GROQ_MODEL_FAST = os.getenv("GROQ_MODEL_FAST", "llama-3.1-8b-instant")

# --- Cerebras (free tier, OpenAI-compatible, even faster) ----------------
CEREBRAS_API_KEY = os.getenv("CEREBRAS_API_KEY", "").strip()
CEREBRAS_BASE_URL = os.getenv("CEREBRAS_BASE_URL", "https://api.cerebras.ai/v1").rstrip("/")
CEREBRAS_MODEL_ANALYZE = os.getenv("CEREBRAS_MODEL_ANALYZE", "llama-3.3-70b")
CEREBRAS_MODEL_DEFAULT = os.getenv("CEREBRAS_MODEL_DEFAULT", "llama-3.3-70b")
CEREBRAS_MODEL_FAST = os.getenv("CEREBRAS_MODEL_FAST", "llama3.1-8b")

# --- GitHub Models (free for devs, OpenAI-compatible) --------------------
GITHUB_MODELS_TOKEN = os.getenv("GITHUB_MODELS_TOKEN", "").strip()
GITHUB_MODELS_BASE_URL = os.getenv(
    "GITHUB_MODELS_BASE_URL", "https://models.inference.ai.azure.com"
).rstrip("/")
GITHUB_MODELS_MODEL_ANALYZE = os.getenv("GITHUB_MODELS_MODEL_ANALYZE", "gpt-4o-mini")
GITHUB_MODELS_MODEL_DEFAULT = os.getenv("GITHUB_MODELS_MODEL_DEFAULT", "gpt-4o-mini")

# --- Ollama (local, offline lifeline) ------------------------------------
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://127.0.0.1:11434").rstrip("/")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3:latest")
OLLAMA_MODEL_ANALYZE = os.getenv("OLLAMA_MODEL_ANALYZE", OLLAMA_MODEL)
OLLAMA_NUM_CTX = int(os.getenv("OLLAMA_NUM_CTX", "2048"))
OLLAMA_NUM_THREAD = int(os.getenv("OLLAMA_NUM_THREAD", "2"))
OLLAMA_NUM_GPU = int(os.getenv("OLLAMA_NUM_GPU", "0"))
OLLAMA_TEMPERATURE = float(os.getenv("OLLAMA_TEMPERATURE", "0.4"))
OLLAMA_NUM_PREDICT_ANALYZE = int(os.getenv("OLLAMA_NUM_PREDICT_ANALYZE", "900"))
OLLAMA_NUM_PREDICT_DEFAULT = int(os.getenv("OLLAMA_NUM_PREDICT_DEFAULT", "420"))
OLLAMA_KEEP_ALIVE = os.getenv("OLLAMA_KEEP_ALIVE", "0m")

# --- Per-task timeouts (seconds) -----------------------------------------
# Cloud LLMs (groq/cerebras/gemini/github) finish in 1-5s; ollama needs much more.
AI_TIMEOUT_ANALYZE = int(os.getenv("AI_TIMEOUT_ANALYZE", "30"))
AI_TIMEOUT_ANALYZE_OLLAMA = int(os.getenv("AI_TIMEOUT_ANALYZE_OLLAMA", "150"))
AI_TIMEOUT_DRAFT = int(os.getenv("AI_TIMEOUT_DRAFT", "30"))
AI_TIMEOUT_ASK = int(os.getenv("AI_TIMEOUT_ASK", "25"))
AI_TIMEOUT_NOTE = int(os.getenv("AI_TIMEOUT_NOTE", "20"))
AI_TIMEOUT_COVER = int(os.getenv("AI_TIMEOUT_COVER", "40"))

# --- Task → provider chain (ordered, comma-separated) --------------------
# A task walks the chain top-to-bottom and uses the first provider that succeeds.
# Providers without an API key are silently skipped, so you can add keys progressively.
def _parse_chain(env_value, default_chain):
    raw = (env_value or "").strip()
    if not raw:
        return list(default_chain)
    parts = [p.strip().lower() for p in raw.split(",") if p.strip()]
    return parts or list(default_chain)

_DEFAULT_CHAIN = ["groq", "cerebras", "gemini", "github", "ollama"]

TASK_CHAIN_ANALYZE = _parse_chain(os.getenv("TASK_CHAIN_ANALYZE"), _DEFAULT_CHAIN)
TASK_CHAIN_DRAFT = _parse_chain(os.getenv("TASK_CHAIN_DRAFT"), _DEFAULT_CHAIN)
TASK_CHAIN_NOTE = _parse_chain(os.getenv("TASK_CHAIN_NOTE"), _DEFAULT_CHAIN)
TASK_CHAIN_ASK = _parse_chain(os.getenv("TASK_CHAIN_ASK"), _DEFAULT_CHAIN)
TASK_CHAIN_COVER = _parse_chain(os.getenv("TASK_CHAIN_COVER"), _DEFAULT_CHAIN)
# Optional: override provider chain for the resume-keyword experiment (defaults to analyze chain).
TASK_CHAIN_RESUME_KEYWORDS = _parse_chain(os.getenv("TASK_CHAIN_RESUME_KEYWORDS"), TASK_CHAIN_ANALYZE)


def _resolve_runtime_path(path, fallback):
    raw = (path or fallback or "").strip() or fallback
    if os.path.isabs(raw):
        return raw
    return os.path.join(PROJECT_ROOT, raw)

# Backward-compatible single-provider env vars (used by quota-health display only).
TASK_PROVIDER_ANALYZE_PRIMARY = TASK_CHAIN_ANALYZE[0] if TASK_CHAIN_ANALYZE else "groq"
TASK_PROVIDER_NOTE = TASK_CHAIN_NOTE[0] if TASK_CHAIN_NOTE else "groq"
TASK_PROVIDER_ASK = TASK_CHAIN_ASK[0] if TASK_CHAIN_ASK else "groq"
TASK_PROVIDER_DRAFT = TASK_CHAIN_DRAFT[0] if TASK_CHAIN_DRAFT else "groq"
TASK_PROVIDER_COVER = TASK_CHAIN_COVER[0] if TASK_CHAIN_COVER else "groq"
USAGE_LOG_PATH = _resolve_runtime_path(os.getenv("AI_USAGE_LOG_PATH"), "data/ai_usage_log.csv")
JOB_TRACKER_PATH = _resolve_runtime_path(os.getenv("JOB_TRACKER_PATH"), "data/job_tracker.csv")
RESUME_DIR = os.getenv("RESUME_DIR", "resumes")
DAILY_ANALYZE_SOFT_CAP = APP_SETTINGS.limit("daily_analyze_soft_cap",
    int(os.getenv("DAILY_ANALYZE_SOFT_CAP", "30")))
DAILY_ANALYZE_HARD_CAP = APP_SETTINGS.limit("daily_analyze_hard_cap",
    int(os.getenv("DAILY_ANALYZE_HARD_CAP", "45")))
SCHEDULED_EMAILS_FILE = _resolve_runtime_path(os.getenv("SCHEDULED_EMAILS_FILE"), "data/scheduled_emails.json")
EMAIL_SCHEDULER_POLL_SECONDS = max(5, APP_SETTINGS.scheduler_int("poll_seconds",
    int(os.getenv("EMAIL_SCHEDULER_POLL_SECONDS", "20"))))
EMAIL_SCHEDULER_HOUR_LOCAL = APP_SETTINGS.scheduler_int("send_at_hour_local",
    int(os.getenv("EMAIL_SCHEDULER_HOUR_LOCAL", "8")))


def _ensure_parent_dir(path):
    """Auto-create parent directory for runtime files so first-run users do not crash."""
    parent = os.path.dirname(path or "")
    if parent and not os.path.isdir(parent):
        try:
            os.makedirs(parent, exist_ok=True)
        except OSError:
            pass


for _runtime_path in (USAGE_LOG_PATH, JOB_TRACKER_PATH, SCHEDULED_EMAILS_FILE):
    _ensure_parent_dir(_runtime_path)

_SCHEDULED_EMAILS_LOCK = threading.Lock()
_scheduler_thread_started = False

# Resume registry is loaded from config/resume_personas.json (or its .example).
# RESUMES   = {code: relative-path-to-pdf-on-disk}
# RESUME_LABELS = {code: human-readable label shown in the UI dropdowns}
RESUMES = PERSONAS.resumes_map
RESUME_LABELS = PERSONAS.labels_map
DEFAULT_RESUME_CODE = PERSONAS.default_code if PERSONAS.default_code in RESUMES else (
    next(iter(RESUMES.keys()), "DA")
)

# ----------------------------------------------------------------------
# CANDIDATE BIO + RESUME PERSONAS — sourced from config files at startup.
# Edit config/profile.json and config/resume_personas.json to change these
# without touching app.py.
# ----------------------------------------------------------------------
RESUME_CONTEXT = (
    "CANDIDATE BIO (true for ALL personas — always available to cite):\n"
    + PROFILE.render_candidate_bio_block()
    + "\n\n"
    + "=========== RESUME PERSONAS — pick ONE per JD ===========\n\n"
    + PERSONAS.render_personas_block()
)


def _load_scheduled_emails():
    if not os.path.exists(SCHEDULED_EMAILS_FILE):
        return []
    try:
        with open(SCHEDULED_EMAILS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_scheduled_emails(items):
    with open(SCHEDULED_EMAILS_FILE, "w", encoding="utf-8") as f:
        json.dump(items, f, indent=2)


def _next_day_send_time_local(now=None):
    """Next-day send time at the configured local hour (default 8:00 AM)."""
    now = now or datetime.now()
    next_day = now.date() + timedelta(days=1)
    hour = max(0, min(23, EMAIL_SCHEDULER_HOUR_LOCAL))
    return datetime.combine(next_day, dt_time(hour=hour, minute=0, second=0, microsecond=0))


# Legacy alias retained for any external callers / older imports.
_next_day_8am_local = _next_day_send_time_local


def _queue_scheduled_email(payload):
    with _SCHEDULED_EMAILS_LOCK:
        jobs = _load_scheduled_emails()
        jobs.append(payload)
        _save_scheduled_emails(jobs)


def _scheduled_email_worker():
    while True:
        try:
            now = datetime.now()
            changed = False
            with _SCHEDULED_EMAILS_LOCK:
                jobs = _load_scheduled_emails()
                for job in jobs:
                    if job.get("status") != "pending":
                        continue
                    send_at_raw = job.get("send_at")
                    if not send_at_raw:
                        continue
                    try:
                        send_at_dt = datetime.fromisoformat(send_at_raw)
                    except ValueError:
                        job["status"] = "failed"
                        job["error"] = "Invalid send_at format"
                        changed = True
                        continue
                    if send_at_dt > now:
                        continue
                    try:
                        send_email(
                            to_email=job["email"],
                            subject=job["subject"],
                            body=job["body"],
                            attachment_path=job["attachment_path"],
                        )
                        job["status"] = "sent"
                        job["sent_at"] = datetime.now().isoformat(timespec="seconds")
                        changed = True
                    except Exception as send_err:
                        job["status"] = "failed"
                        job["error"] = str(send_err)
                        changed = True
                if changed:
                    _save_scheduled_emails(jobs)
        except Exception as worker_err:
            print(f"[scheduler] Scheduled email worker error: {worker_err}")
        time.sleep(EMAIL_SCHEDULER_POLL_SECONDS)


def _ensure_scheduler_started():
    global _scheduler_thread_started
    if _scheduler_thread_started:
        return
    with _SCHEDULED_EMAILS_LOCK:
        if _scheduler_thread_started:
            return
        t = threading.Thread(target=_scheduled_email_worker, daemon=True)
        t.start()
        _scheduler_thread_started = True


_ensure_scheduler_started()

USER_CONTEXT = (
    PROFILE.render_user_context_block()
    + "\n\n"
    + "EXPERIENCE BASELINE (use this for every role—do not ignore it):\n"
    + "- When the JD states a minimum years-of-experience requirement (e.g. '5+ years', '8+ years'), compare that requirement explicitly to the baseline above. If the JD asks for materially more years than this profile, you MUST flag it—typically RED LIGHT for clear multi-year gaps, YELLOW when borderline or when other signals (exceptional projects) might partially offset.\n"
    + "- A strong technical stack match does NOT erase a seniority gap; keep those assessments separate.\n"
)

@app.route("/")
def index():
    return render_template(
        "index.html",
        enable_followup_agent=ENABLE_FOLLOWUP_AGENT,
        enable_email_sending=ENABLE_EMAIL_SENDING,
        enable_email_scheduling=ENABLE_EMAIL_SCHEDULING,
        app_title=APP_SETTINGS.ui_text("app_title", "Career Command Center"),
        app_subtitle=APP_SETTINGS.ui_text("app_subtitle", "AI-Powered Job Fit & Outreach Engine"),
        scheduler_hour_label=f"{EMAIL_SCHEDULER_HOUR_LOCAL:02d}:00",
        alumni_search_enabled=PROFILE.alumni_search_enabled and APP_SETTINGS.feature("enable_alumni_search_button", True),
        alumni_button_label=PROFILE.alumni_button_label,
        alumni_button_emoji=PROFILE.alumni_button_emoji,
        alumni_school_slug=PROFILE.alumni_school_slug,
    )


def _load_recent_tracker_rows(limit=40):
    tracker_path = JOB_TRACKER_PATH
    if not os.path.exists(tracker_path):
        return []
    rows = []
    with open(tracker_path, mode="r", encoding="utf-8", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            rows.append(row)
    return rows[-limit:]


def _build_tracker_context(rows):
    if not rows:
        return "No tracker entries are available yet."
    lines = []
    for i, row in enumerate(rows, start=1):
        lines.append(
            f"{i}. Date={row.get('Date', 'N/A')}; Company={row.get('Company', 'N/A')}; "
            f"Position={row.get('Position', 'N/A')}; Resume={row.get('Resume Used', 'N/A')}; "
            f"Platform={row.get('Platform', 'N/A')}; Status={row.get('Status', 'N/A')}; "
            f"Contact={row.get('Contact Name', 'N/A')}; Message={row.get('Message', 'N/A')}"
        )
    return "\n".join(lines)


def _estimate_tokens(text):
    # Lightweight estimate to monitor usage across providers consistently.
    return max(1, len((text or "")) // 4)


def _log_ai_usage(endpoint, provider, model, prompt_text, response_text, latency_ms, status="ok", error_code=""):
    file_exists = os.path.exists(USAGE_LOG_PATH)
    with open(USAGE_LOG_PATH, mode="a", encoding="utf-8", newline="") as file:
        writer = csv.writer(file)
        if not file_exists:
            writer.writerow([
                "timestamp",
                "endpoint",
                "provider",
                "model",
                "prompt_chars",
                "response_chars",
                "prompt_tokens_est",
                "response_tokens_est",
                "latency_ms",
                "status",
                "error_code",
            ])
        writer.writerow([
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            endpoint,
            provider,
            model,
            len(prompt_text or ""),
            len(response_text or ""),
            _estimate_tokens(prompt_text),
            _estimate_tokens(response_text),
            int(latency_ms),
            status,
            error_code,
        ])


def _read_usage_summary_today():
    base = {
        "analyze_today": 0,
        "quota_429_today": 0,
        "requests_today": 0,
        "gemini_today": 0,
        "ollama_today": 0,
        "groq_today": 0,
        "cerebras_today": 0,
        "github_today": 0,
    }
    if not os.path.exists(USAGE_LOG_PATH):
        return base
    today = datetime.now().strftime("%Y-%m-%d")
    stats = dict(base)
    with open(USAGE_LOG_PATH, mode="r", encoding="utf-8", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            ts = row.get("timestamp", "")
            if not ts.startswith(today):
                continue
            stats["requests_today"] += 1
            provider = (row.get("provider") or "").strip().lower()
            key = f"{provider}_today"
            if key in stats:
                stats[key] += 1
            if row.get("endpoint") == "/api/analyze":
                stats["analyze_today"] += 1
            err = (row.get("error_code") or "").lower()
            if "429" in err or "quota" in err or "resource_exhausted" in err:
                stats["quota_429_today"] += 1
    return stats


def _should_fallback_to_ollama(err):
    """Legacy helper kept for the old analyze fallback path. Always advance on these failures."""
    raw = str(err)
    upper = raw.upper()
    return (
        "429" in raw
        or "QUOTA" in upper
        or "RESOURCE_EXHAUSTED" in upper
        or "RATE LIMIT" in upper
        or "TIMED OUT" in upper
        or "TIMEOUT" in upper
    )


def _provider_is_available(provider):
    """Skip a provider in the chain if its credentials aren't configured.

    This lets the user add API keys progressively without breaking the workflow.
    """
    p = (provider or "").strip().lower()
    if p == "groq":
        return bool(GROQ_API_KEY)
    if p == "cerebras":
        return bool(CEREBRAS_API_KEY)
    if p == "github":
        return bool(GITHUB_MODELS_TOKEN)
    if p == "gemini":
        return bool(api_key)
    if p == "ollama":
        return True  # local; cheap to attempt, fails fast if Ollama isn't running
    return False


def _provider_timeout_for_task(provider, task):
    """Cloud providers respond in <5s; only Ollama needs the long window."""
    p = (provider or "").strip().lower()
    if p == "ollama":
        if task == "analyze":
            return AI_TIMEOUT_ANALYZE_OLLAMA
        if task == "cover":
            return max(AI_TIMEOUT_COVER, 90)
        if task == "draft":
            return max(AI_TIMEOUT_DRAFT, 75)
        if task == "note":
            return max(AI_TIMEOUT_NOTE, 60)
        if task == "ask":
            return max(AI_TIMEOUT_ASK, 60)
        return 90
    if task == "analyze":
        return AI_TIMEOUT_ANALYZE
    if task == "cover":
        return AI_TIMEOUT_COVER
    if task == "draft":
        return AI_TIMEOUT_DRAFT
    if task == "note":
        return AI_TIMEOUT_NOTE
    if task == "ask":
        return AI_TIMEOUT_ASK
    if task == "resume_keywords":
        if p == "ollama":
            return AI_TIMEOUT_ANALYZE_OLLAMA
        return AI_TIMEOUT_ANALYZE
    return 30


# =====================================================================
# PROVIDER CIRCUIT BREAKER
# =====================================================================
# Without a breaker, a broken provider (e.g. GitHub Models with a PAT
# missing the `models:read` scope, or Cerebras with a wrong model name)
# gets retried on EVERY analyze/note/draft request — wasting ~300ms to
# 3s of user-visible latency every time, because the chain has to walk
# through the corpse before reaching a working provider.
#
# The breaker tracks per-provider health in-process and short-circuits
# doomed calls. Rules:
#
#   - Permanent bad-credential errors (401, 403, 404):
#       The provider is poisoned for the entire server lifetime — no
#       amount of waiting fixes a wrong PAT scope or wrong model name.
#       The user needs to fix env + restart. We skip those providers on
#       ALL subsequent calls until restart.
#
#   - Rate-limit / quota errors (429):
#       The provider self-heals after a cooldown. We park it for
#       PROVIDER_QUOTA_COOLDOWN_SECONDS (default 10 minutes), then retry.
#
#   - Timeouts (Ollama spinning on a big prompt, or a flaky cloud call):
#       We count consecutive timeouts. After 3 in a row, we park the
#       provider for PROVIDER_TIMEOUT_COOLDOWN_SECONDS (default 5 min).
#       A successful call resets the counter.
#
# Successful calls ALWAYS clear that provider's fault record.
# =====================================================================

_PROVIDER_HEALTH = {}  # provider -> {"dead_until": float | "forever", "reason": str, "timeouts": int}
PROVIDER_QUOTA_COOLDOWN_SECONDS = 10 * 60     # 10 min park on generic 429
PROVIDER_DAILY_QUOTA_COOLDOWN_SECONDS = 12 * 3600  # 12 hr park when the 429 is a DAILY cap
PROVIDER_TIMEOUT_COOLDOWN_SECONDS = 5 * 60    # 5 min park after 3 timeouts
PROVIDER_MAX_CONSECUTIVE_TIMEOUTS = 3
# Wall-clock budget for a full provider-chain walk. If providers 1..N-1
# have already burned this much combined, we stop and return the last
# error rather than waiting for the slow local model to grind. This is
# what prevents the "40s elapsed then Ollama adds another 150s" scenario.
CHAIN_WALL_CLOCK_BUDGET_SECONDS = 45.0


def _provider_is_parked(provider):
    """Returns (is_parked, reason_string). Used by the chain to skip
    providers that are known-dead without paying the failure cost."""
    entry = _PROVIDER_HEALTH.get(provider)
    if not entry:
        return False, None
    dead_until = entry.get("dead_until")
    if dead_until == "forever":
        return True, entry.get("reason") or "permanently parked"
    if isinstance(dead_until, (int, float)) and time.time() < dead_until:
        remaining = int(dead_until - time.time())
        return True, f"{entry.get('reason') or 'cooling down'} ({remaining}s left)"
    # Cooldown elapsed -> clear the park so the next call retries.
    if dead_until:
        _PROVIDER_HEALTH.pop(provider, None)
    return False, None


def _classify_provider_error(err):
    """Inspect an exception and decide how the breaker should react.
    Returns one of: 'forever' (unrecoverable), 'quota_daily' (park 12h),
    'quota' (park 10min), 'timeout' (count consecutive), 'soft' (ignore)."""
    msg = str(err or "").lower()
    # Timeouts.
    if isinstance(err, TimeoutError) or "timed out" in msg or "timeout" in msg:
        return "timeout"
    # Cerebras queue overflow ("We're experiencing high traffic right now! ...
    # too_many_requests_error / queue_exceeded"). This is short-term capacity,
    # NOT a credentials problem — treat as quota so we retry in 10 min, not
    # mark forever-dead.
    if "queue_exceeded" in msg or "too_many_requests_error" in msg or "high traffic" in msg:
        return "quota"
    # DAILY-cap rate-limits: these won't self-heal in 10 minutes, so park hard.
    # Groq uses "tokens per day (TPD)" and a retry delay on the order of hours.
    # Gemini free tier says "free_tier_requests" / "free_tier_input_token_count"
    # with limit: 0 when the day is burned.
    if ("429" in msg or "quota" in msg or "rate limit" in msg) and any(tag in msg for tag in (
        "per day", "tokens per day", "tpd", "free_tier", "free tier",
        "daily", "limit: 0", "requests per day", "rpd",
    )):
        return "quota_daily"
    # Generic rate-limit / quota (per-minute etc, self-heals shortly).
    if "429" in msg or "rate limit" in msg or "quota" in msg or "resourceexhausted" in msg or "resource_exhausted" in msg:
        return "quota"
    # Permanent auth/availability errors.
    if any(code in msg for code in ("http 401", "http 403", "http 404")):
        return "forever"
    if "unauthorized" in msg or "permission" in msg or "does not exist" in msg or "not_found" in msg or "model_not_found" in msg:
        return "forever"
    # Un-parseable JSON = model weirdness, not provider death.
    if "un-parseable json" in msg or "invalid json" in msg:
        return "soft"
    return "soft"


def _mark_provider_failed(provider, err):
    """Record a failure. Returns the classification so the caller can log it."""
    kind = _classify_provider_error(err)
    entry = _PROVIDER_HEALTH.setdefault(provider, {"timeouts": 0})
    if kind == "forever":
        entry["dead_until"] = "forever"
        entry["reason"] = f"unrecoverable: {str(err)[:120]}"
    elif kind == "quota_daily":
        entry["dead_until"] = time.time() + PROVIDER_DAILY_QUOTA_COOLDOWN_SECONDS
        entry["reason"] = "daily quota exhausted (parked 12h)"
    elif kind == "quota":
        entry["dead_until"] = time.time() + PROVIDER_QUOTA_COOLDOWN_SECONDS
        entry["reason"] = "quota/rate-limit"
    elif kind == "timeout":
        entry["timeouts"] = entry.get("timeouts", 0) + 1
        if entry["timeouts"] >= PROVIDER_MAX_CONSECUTIVE_TIMEOUTS:
            entry["dead_until"] = time.time() + PROVIDER_TIMEOUT_COOLDOWN_SECONDS
            entry["reason"] = f"{entry['timeouts']} consecutive timeouts"
    # 'soft' failures don't poison the breaker — just clear so a success after
    # them resets cleanly.
    return kind


def _mark_provider_success(provider):
    """Successful call clears any fault record for this provider."""
    _PROVIDER_HEALTH.pop(provider, None)


@app.route("/api/provider-health", methods=["GET", "POST"])
def provider_health():
    """Expose the breaker state so the frontend can show which providers
    are currently parked — and let the user force-reset them without a
    server restart (useful after fixing a PAT scope or model name)."""
    if request.method == "POST":
        _PROVIDER_HEALTH.clear()
        return jsonify({"reset": True})
    out = {}
    now = time.time()
    for provider, entry in _PROVIDER_HEALTH.items():
        dead_until = entry.get("dead_until")
        if dead_until == "forever":
            out[provider] = {"state": "dead", "reason": entry.get("reason"), "remaining_s": None}
        elif isinstance(dead_until, (int, float)):
            remaining = max(0, int(dead_until - now))
            out[provider] = {"state": "parked", "reason": entry.get("reason"), "remaining_s": remaining}
        else:
            out[provider] = {"state": "degraded", "reason": entry.get("reason"), "timeouts": entry.get("timeouts", 0)}
    return jsonify({"providers": out})


def _run_provider_chain(task, chain, prompt_builder, expect_json=True, on_error_log=None):
    """Walk a provider chain and return the first successful response.

    prompt_builder(provider) -> str  lets us tune the prompt per provider
    (e.g. compress JD aggressively for Ollama, send full JD to Groq).
    When expect_json=True, a provider that returns un-parseable JSON is
    treated as a soft failure so the chain advances to the next provider
    instead of bubbling a parse error back to the user.
    Returns: (response_text, model_used, provider_used, prompt_used, attempted)
    """
    last_err = None
    attempted = []
    chain_started = time.time()
    for provider in chain:
        provider = (provider or "").strip().lower()
        if not provider:
            continue
        if not _provider_is_available(provider):
            attempted.append(f"{provider}(skip:no-key)")
            continue
        # Skip providers the breaker has parked — this is the fast-path that
        # eliminates the multi-second "walk through 4 dead providers" stall.
        is_parked, park_reason = _provider_is_parked(provider)
        if is_parked:
            attempted.append(f"{provider}(skip:{park_reason})")
            continue
        # Wall-clock guard: if previous providers have already burned the
        # budget, don't launch the slow local model — a user who's been
        # waiting 45s doesn't want another 150s added on. Cloud providers
        # are cheap enough to always try, so we only skip ollama when
        # over-budget.
        elapsed_so_far = time.time() - chain_started
        if elapsed_so_far > CHAIN_WALL_CLOCK_BUDGET_SECONDS and provider == "ollama":
            attempted.append(f"{provider}(skip:chain-budget-{int(elapsed_so_far)}s)")
            continue
        prompt = prompt_builder(provider)
        timeout = _provider_timeout_for_task(provider, task)
        attempt_started = time.time()
        try:
            text, model_used = _ai_generate_text_with_provider(
                prompt, provider, timeout_seconds=timeout, expect_json=expect_json, task=task,
            )
            # If we promised the caller JSON, validate it here so a malformed
            # response from one provider doesn't reach the route. The chain
            # then tries the next provider just like for any other error.
            if expect_json:
                try:
                    json.loads(_strip_markdown_fences(text))
                except Exception as parse_err:
                    raise Exception(
                        f"{provider} returned un-parseable JSON: {parse_err}"
                    ) from parse_err
            _mark_provider_success(provider)
            attempted.append(f"{provider}(ok)")
            return text, model_used, provider, prompt, attempted
        except Exception as e:
            last_err = e
            kind = _mark_provider_failed(provider, e)
            attempted.append(f"{provider}(err:{kind})")
            if on_error_log is not None:
                try:
                    on_error_log(provider, prompt, e, (time.time() - attempt_started) * 1000)
                except Exception:
                    pass
            # All cross-provider errors warrant trying the next provider:
            # auth issue, quota, timeout, network, malformed JSON, model-unavailable
            # -- the next provider in the chain has independent credentials and infra.
            continue
    # When the WHOLE chain fails, raise a single summary exception rather
    # than whatever garbled last provider error happened to be in flight.
    # This is what the /api/analyze route's error handler converts into a
    # friendly alert. Including `attempted` helps the user understand which
    # providers failed and why (quota, parked, 401, etc).
    summary = (
        f"All providers failed or unavailable for task={task}. "
        f"Tried: {', '.join(attempted) or 'none'}."
    )
    if last_err:
        summary += f" Last error from chain: {str(last_err)[:240]}"
    raise Exception(summary)


def _http_post_json(url, headers, payload, timeout_seconds):
    """Tiny urllib JSON POST helper used by all OpenAI-compatible providers."""
    body = json.dumps(payload).encode("utf-8")
    # Cloudflare (used by Groq + Cerebras) blocks default Python-urllib User-Agents
    # with "error code 1010". Force a real-looking UA so the request is accepted.
    full_headers = {
        "User-Agent": "career-command-center/1.0 (+python; macOS)",
        "Accept": "application/json",
    }
    full_headers.update(headers or {})
    req = url_request.Request(url, data=body, headers=full_headers, method="POST")
    # Some Python installs on macOS don't trust the system keychain by default,
    # which breaks HTTPS calls with CERTIFICATE_VERIFY_FAILED. Force certifi's
    # CA bundle so Groq/Cerebras/GitHub requests verify consistently.
    ssl_ctx = ssl.create_default_context(cafile=certifi.where())
    try:
        with url_request.urlopen(req, timeout=timeout_seconds, context=ssl_ctx) as resp:
            return resp.read().decode("utf-8")
    except url_error.HTTPError as e:
        details = e.read().decode("utf-8", errors="replace") if hasattr(e, "read") else str(e)
        # Surface the upstream error code so _should_fallback_to_ollama and the chain can react.
        raise Exception(f"HTTP {e.code}: {details}") from e
    except url_error.URLError as e:
        raise Exception(f"Network error contacting {url}: {e}") from e
    except TimeoutError as e:
        raise TimeoutError(f"AI request timed out after {timeout_seconds} seconds") from e


def _call_openai_compatible(base_url, api_key_value, model_name, prompt,
                            timeout_seconds=30, expect_json=False,
                            temperature=0.3, max_tokens=1500,
                            extra_headers=None):
    """OpenAI Chat Completions wire format used by Groq, Cerebras, and GitHub Models."""
    headers = {
        "Authorization": f"Bearer {api_key_value}",
        "Content-Type": "application/json",
    }
    if extra_headers:
        headers.update(extra_headers)
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False,
    }
    if expect_json:
        # Groq, Cerebras, and gpt-4o-mini all honor this hint; non-supporting providers
        # ignore it harmlessly. Combined with our prompt's "Return ONLY valid JSON" guard,
        # this gets us reliable parseable output.
        payload["response_format"] = {"type": "json_object"}

    body = _http_post_json(f"{base_url}/chat/completions", headers, payload, timeout_seconds)
    parsed = json.loads(body)
    choices = parsed.get("choices") or []
    if not choices:
        raise Exception(f"Empty completion from {base_url}: {body[:200]}")
    message = choices[0].get("message") or {}
    content = message.get("content")
    if not content or not str(content).strip():
        raise Exception(f"Empty content from {base_url} (model={model_name})")
    return str(content).strip()


def _model_for_groq(task):
    if task == "analyze" or task == "cover" or task == "draft":
        return GROQ_MODEL_ANALYZE
    if task == "note" or task == "ask":
        return GROQ_MODEL_FAST
    return GROQ_MODEL_DEFAULT


def _model_for_cerebras(task):
    if task == "analyze" or task == "cover" or task == "draft":
        return CEREBRAS_MODEL_ANALYZE
    if task == "note" or task == "ask":
        return CEREBRAS_MODEL_FAST
    return CEREBRAS_MODEL_DEFAULT


def _model_for_github(task):
    if task == "analyze" or task == "cover" or task == "draft":
        return GITHUB_MODELS_MODEL_ANALYZE
    return GITHUB_MODELS_MODEL_DEFAULT


def _temperature_for_task(task):
    # Lower temp = more deterministic JSON; slightly higher for creative writing.
    if task in ("analyze", "ask"):
        return 0.25
    if task in ("note", "draft", "cover"):
        return 0.55
    return 0.4


def _build_emergency_connection_note(company, position, contact_name):
    _, clean_contact_name = _extract_clean_contact_name(contact_name)
    first_name = clean_contact_name.split(" ")[0] if clean_contact_name else "there"
    company_text = (company or "your company").strip() or "your company"
    role_text = (position or "the role").strip() or "the role"
    note = (
        f"Hi {first_name}, I am exploring the {role_text} opportunity at {company_text}. "
        "My background in analytics and cross-functional execution aligns well with this scope, "
        "and I would value connecting to learn how your team approaches this work."
    )
    return _normalize_connection_note(note)


def _friendly_ai_error(err, default_prefix):
    raw = str(err)
    upper = raw.upper()
    # When the chain walker raises its summary exception, unpack it into a
    # compact human-readable message. This is the most common path after
    # the chain exhausts and we don't want the user to see raw 401 JSON.
    if "ALL PROVIDERS FAILED" in upper:
        reasons = []
        # attempted is embedded as "Tried: groq(err:quota_daily), cerebras(err:forever), ..."
        tried_match = re.search(r"Tried:\s*([^.]+)\.", raw)
        if tried_match:
            for chunk in tried_match.group(1).split(","):
                chunk = chunk.strip()
                if "(skip:daily" in chunk or "(err:quota_daily" in chunk:
                    reasons.append(f"• {chunk.split('(')[0]}: daily quota exhausted (will retry tomorrow)")
                elif "(skip:quota" in chunk or "(err:quota" in chunk:
                    reasons.append(f"• {chunk.split('(')[0]}: rate-limited (retrying soon)")
                elif "(err:forever" in chunk or "(skip:unrecoverable" in chunk:
                    if "github" in chunk.lower():
                        reasons.append(f"• {chunk.split('(')[0]}: PAT is missing the 'models:read' scope — generate a new one at github.com/settings/tokens")
                    elif "cerebras" in chunk.lower():
                        reasons.append(f"• {chunk.split('(')[0]}: model name invalid (check CEREBRAS_MODEL_* in .env)")
                    else:
                        reasons.append(f"• {chunk.split('(')[0]}: credentials rejected (check .env)")
                elif "(skip:no-key" in chunk:
                    reasons.append(f"• {chunk.split('(')[0]}: no API key configured")
                elif "(err:timeout" in chunk:
                    reasons.append(f"• {chunk.split('(')[0]}: timed out")
                elif "(skip:chain-budget" in chunk:
                    reasons.append(f"• {chunk.split('(')[0]}: skipped (chain already over budget)")
        joined = "\n".join(reasons) if reasons else "(see server log for full chain details)"
        return (
            "All AI providers are currently unavailable:\n\n"
            f"{joined}\n\n"
            "What to do right now:\n"
            "• Click 'Skip to Outreach Strategy' to send a draft manually.\n"
            "• Or wait for quotas to reset (~tomorrow for daily caps).\n"
            "• Or fix the broken credentials in .env and restart.",
            503,
        )
    if "OLLAMA" in upper and ("REFUSED" in upper or "UNREACHABLE" in upper or "NOT RUNNING" in upper):
        return (
            "Ollama is not reachable. Start it with 'ollama serve' and ensure OLLAMA_URL is correct.",
            503,
        )
    if "MODEL" in upper and "NOT FOUND" in upper:
        return (
            "Configured model was not found. Pull it with 'ollama pull <model>' or update the model name in .env.",
            400,
        )
    if ("HTTP 401" in upper) or ("UNAUTHORIZED" in upper):
        return (
            "An AI provider rejected the API key (HTTP 401). Re-check the key in .env. Details: " + raw,
            401,
        )
    if "TIMED OUT" in upper or "TIMEOUT" in upper:
        return (
            "The AI request took too long and was stopped. Please try again in a few seconds.",
            504,
        )
    if "429" in raw or "QUOTA" in upper or "RATE LIMIT" in upper or "RESOURCE_EXHAUSTED" in upper:
        retry_match = re.search(r"retry in ([0-9]+(?:\.[0-9]+)?)s", raw, flags=re.IGNORECASE)
        retry_hint = ""
        if retry_match:
            retry_hint = f" Please retry in about {int(float(retry_match.group(1)))} seconds."
        msg = (
            "An AI provider returned quota/rate-limit. The chain advanced but every "
            "provider was exhausted or unavailable."
            f"{retry_hint} Details: {raw}"
        )
        return msg, 429
    return f"{default_prefix}: {raw}", 500


def _build_discovery_input(payload):
    source_url = (payload.get("source_url") or "").strip()
    company = (payload.get("company") or "").strip()
    position = (payload.get("position") or "").strip()
    jd = (payload.get("jd") or "").strip()
    if payload.get("actor_input") and isinstance(payload.get("actor_input"), dict):
        return payload["actor_input"]
    return {
        "sourceUrl": source_url,
        "company": company,
        "position": position,
        "jd": jd,
        "startUrls": ([{"url": source_url}] if source_url else []),
    }


def _extract_enrichment_handoff(items):
    """Best-effort mapper from actor output to existing form fields."""
    if not items:
        return {}
    first = items[0] if isinstance(items[0], dict) else {}
    contact_name = first.get("contact_name") or first.get("fullName") or first.get("name") or ""
    contact_role = first.get("contact_role") or first.get("title") or first.get("role") or ""
    target_email = first.get("target_email") or first.get("email") or ""
    company = first.get("company") or ""
    position = first.get("position") or ""
    out = {
        "company": str(company).strip(),
        "position": str(position).strip(),
        "contact_name": str(contact_name).strip(),
        "contact_role": str(contact_role).strip(),
        "target_email": str(target_email).strip(),
    }
    return {k: v for k, v in out.items() if v}


def _normalize_apify_error(err):
    raw = str(err)
    upper = raw.upper()
    if "APIFY_TOKEN" in upper:
        return "Apify is not configured (missing APIFY_TOKEN).", 503
    if "ACTOR" in upper and "REQUIRED" in upper:
        return "Apify actor IDs are not configured in .env.", 503
    if "FAILED AFTER" in upper:
        return f"Apify actor run failed after retries: {raw}", 502
    return f"Apify request failed: {raw}", 500


def _parse_optional_json_env(name: str) -> dict:
    raw = os.getenv(name, "").strip()
    if not raw:
        return {}
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _deep_merge_dicts(base: dict, extra: dict) -> dict:
    out = dict(base or {})
    for key, val in (extra or {}).items():
        if isinstance(val, dict) and isinstance(out.get(key), dict):
            out[key] = _deep_merge_dicts(out[key], val)
        else:
            out[key] = val
    return out


def _build_linkedin_jobs_search_url(keyword: str, posted_within_seconds: int) -> str:
    kw = (keyword or "").strip()
    secs = max(1, int(posted_within_seconds))
    encoded = quote(kw, safe="")
    # LinkedIn encodes \"Date posted\" as f_TPR=r<seconds>; r86400 = past 24h.
    url = (
        "https://www.linkedin.com/jobs/search/?keywords="
        f"{encoded}&sortBy=R&f_TPR=r{secs}"
    )
    # Optional: align with a manual search (e.g. United States geoId=103644278).
    geo = (os.getenv("SCOUT_LINKEDIN_GEO_ID") or "").strip()
    if geo:
        sep = "&" if "?" in url else "?"
        url = f"{url}{sep}geoId={quote(geo, safe='')}"
    return url


def _linkedin_jobs_actor_input(
    *,
    linkedin_url: str,
    keyword: str,
    max_items: int,
    actor_id: str,
    extras: dict,
) -> dict:
    """
    Map our scout params to Apify actor input.

    curious_coder/linkedin-jobs-scraper expects:
      { \"urls\": [\"https://...\"], \"count\": N }  (count minimum 10 in schema)

    We previously sent urls as [{\"url\": ...}] which yields empty datasets for that actor.
    """
    count = max(10, min(int(max_items), 1000))
    actor_lower = (actor_id or "").lower()

    # curious_coder linkedin-jobs-scraper (public search URL only)
    if "linkedin-jobs-scraper" in actor_lower and "unlimited" not in actor_lower:
        base = {
            "urls": [linkedin_url],
            "count": count,
        }
    else:
        # Broad fallback for other actors (Puppeteer startUrls style, etc.)
        base = {
            "startUrls": [{"url": linkedin_url}],
            "urls": [linkedin_url],
            "searchUrl": linkedin_url,
            "keyword": keyword,
            "keywords": keyword,
            "maxItems": max_items,
        }

    merged = _deep_merge_dicts(base, extras or {})

    # Normalize urls: some templates use [{ "url": "..." }]; this actor needs string[].
    u = merged.get("urls")
    if isinstance(u, list) and u and isinstance(u[0], dict):
        merged["urls"] = [str(x.get("url") or "") for x in u if isinstance(x, dict) and x.get("url")]
    # Map maxItems -> count when count missing (curious_coder schema uses count).
    if "count" not in merged and merged.get("maxItems") is not None:
        try:
            merged["count"] = max(10, min(int(merged["maxItems"]), 1000))
        except Exception:
            pass

    return merged


def _extract_job_candidates_from_items(source: str, items: list) -> list:
    rows = []
    if not isinstance(items, list):
        return rows
    for raw in items:
        if not isinstance(raw, dict):
            continue
        title = raw.get("title") or raw.get("jobTitle") or raw.get("position") or raw.get("name") or ""
        comp_val = raw.get("company") or raw.get("companyName")
        if isinstance(comp_val, dict):
            company = comp_val.get("name") or ""
        else:
            company = comp_val or ""
        url = (
            raw.get("url")
            or raw.get("absoluteUrl")
            or raw.get("absolute_url")
            or raw.get("postingUrl")
            or raw.get("posting_url")
            or raw.get("hostedUrl")
            or raw.get("hosted_url")
            or raw.get("jobUrl")
            or raw.get("jobPostingUrl")
            or raw.get("link")
            or raw.get("applyUrl")
            or raw.get("apply_url")
            or ""
        )
        posted = raw.get("postedAt") or raw.get("postedTime") or raw.get("posted") or ""
        rows.append(
            {
                "title": str(title).strip(),
                "company": str(company).strip(),
                "url": str(url).strip(),
                "source": source,
                "posted_hint": str(posted).strip(),
                "raw": raw,
            }
        )
    return rows


def _merge_jobs_by_url(records: list) -> list:
    seen = set()
    merged = []
    for row in records or []:
        url = str(row.get("url") or "").strip()
        title = str(row.get("title") or "").strip()
        company = str(row.get("company") or "").strip()
        key = url if url else f"{company}|{title}"
        norm = key.lower()
        if not norm:
            continue
        if norm in seen:
            continue
        seen.add(norm)
        item = dict(row)
        raw = item.get("raw")
        if isinstance(raw, dict):
            desc = (
                raw.get("description")
                or raw.get("jobDescription")
                or raw.get("descriptionText")
                or raw.get("description_html")
                or raw.get("jobDescriptionHtml")
            )
            if desc:
                item["description"] = str(desc)
        item.pop("raw", None)
        merged.append(item)
    return merged


def _haystack_for_scout_job(job: dict) -> tuple[str, str]:
    """Full-text haystack plus title-only substring for persona matching (local, fast)."""
    title = str(job.get("title") or "")
    company = str(job.get("company") or "")
    desc = str(job.get("description") or "")[:1400]
    blob = f"{title} {company} {desc}".lower()
    return blob, title.lower()


def _score_scout_job_against_personas(job: dict) -> tuple[float, list[str]]:
    """
    Lightweight overlap score vs resume_personas triggers + core_stack.
    NOT an LLM fit score — ranks listings so likely matches float up without extra API calls.
    """
    hay, title_l = _haystack_for_scout_job(job)
    raw = 0.0
    hints: list[str] = []

    def _note(code: str, phrase: str) -> None:
        tag = f"{code}: matched \"{phrase[:48]}\""
        if len(hints) < 10:
            hints.append(tag)

    for persona in PERSONAS.personas:
        code = str(persona.get("code") or "").strip().upper() or "?"
        for tr in persona.get("triggers") or []:
            t = str(tr).strip().lower()
            if len(t) < 3:
                continue
            if t not in hay:
                continue
            w = min(14.0, 2.5 + len(t) * 0.45)
            if t in title_l:
                w += 5.0
            raw += w
            _note(code, t)
        for stk in persona.get("core_stack") or []:
            s = str(stk).strip().lower()
            if not s:
                continue
            for piece in re.split(r"[,;/]+", s):
                piece = piece.strip().lower()
                if len(piece) < 2:
                    continue
                if piece not in hay:
                    continue
                w = min(10.0, 2.0 + len(piece) * 0.35)
                if piece in title_l:
                    w += 3.0
                raw += w
                _note(code, piece)
        label = str(persona.get("label") or "").strip()
        lbl = label.lower()
        if len(lbl) >= 5 and lbl in hay:
            raw += min(18.0, 5.0 + len(lbl) * 0.2)
            if lbl in title_l:
                raw += 6.0
            _note(code, label[:40])

    return raw, hints


def _rank_scout_jobs_by_personas(merged: list) -> list:
    """Sort by persona keyword overlap (desc); add relevance_score 0–100 within this batch."""
    if not merged:
        return merged
    scored = []
    for row in merged:
        raw_pts, hints = _score_scout_job_against_personas(row)
        scored.append({"row": row, "raw": raw_pts, "hints": hints})
    scored.sort(key=lambda x: (-x["raw"], (x["row"].get("title") or "").lower()))

    raws = [x["raw"] for x in scored]
    hi = max(raws)
    lo = min(raws)

    out: list = []
    for item in scored:
        row = item["row"]
        r = item["raw"]
        hints = item["hints"]
        if hi > lo:
            rel = int(round(100 * (r - lo) / (hi - lo)))
        else:
            rel = 82 if r > 0 else 38
        row["relevance_score"] = max(0, min(100, rel))
        row["relevance_hint"] = "; ".join(hints[:8]) if hints else (
            "Limited overlap with your persona keywords in the listing text — "
            "open the job or run Deep Strategic Analysis for a real fit score."
        )
        out.append(row)
    return out


def _gather_resume_keyword_blocks():
    """Load every persona PDF into text blocks for resume→keyword AI. Returns (blocks, extraction_errors)."""
    try:
        max_per = max(500, int(os.getenv("RESUME_KEYWORD_MAX_CHARS_PER_PDF", "6000")))
    except Exception:
        max_per = 6000
    blocks: list = []
    extraction_errors: list = []
    for code in sorted(PERSONAS.codes):
        rel = RESUMES.get(code)
        if not rel:
            continue
        full = rel if os.path.isabs(rel) else os.path.join(PROJECT_ROOT, rel)
        meta = PERSONAS.get(code) or {}
        label = (meta.get("label") or code).strip()
        try:
            resume_text = extract_pdf_text(full, max_chars=max_per)
        except ResumePdfError as err:
            extraction_errors.append({"code": code, "error": str(err)})
            continue
        blocks.append({
            "code": code,
            "label": label,
            "triggers": meta.get("triggers") or [],
            "core_stack": meta.get("core_stack") or [],
            "resume_text": resume_text,
        })
    return blocks, extraction_errors


def _build_resume_keyword_experiment_prompt(blocks: list) -> str:
    """blocks: list of {code, label, triggers, core_stack, resume_text}"""
    lines = [
        "You are an expert career advisor. The candidate has multiple resume PDF variants (personas).",
        "Each block below includes CONFIG hints (triggers, stack) plus EXTRACTED TEXT from their PDF.",
        "",
        "TASK: For every persona, infer realistic LinkedIn job-search keyword phrases and target job titles.",
        "Also produce a combined list that spans all personas (for broad discovery).",
        "",
        "RULES:",
        "- Return ONLY valid JSON (no markdown fences).",
        '- "linkedin_search_phrases" must be short strings a human would type into LinkedIn Jobs search.',
        '- "primary_linkedin_search" must be ONE concise string (under ~120 chars) that best matches ',
        "  ALL personas together — this will be pasted into LinkedIn Jobs search as-is.",
        "- Do not invent company names, job posting URLs, or employers.",
        "- Ground every suggestion in the resume text; note gaps honestly in \"notes\".",
        "- Prefer inclusive, modern job-market language; avoid limiting the candidate to one narrow niche unless the resume clearly supports it.",
        "",
        "REQUIRED JSON SHAPE:",
        "{",
        '  "primary_linkedin_search": "single best LinkedIn Jobs query for all personas",',
        '  "per_persona": [',
        "    {",
        '      "code": "DA",',
        '      "label": "string",',
        '      "linkedin_search_phrases": ["string", "..."],',
        '      "target_role_titles": ["string", "..."],',
        '      "one_line_focus": "string"',
        "    }",
        "  ],",
        '  "combined_top_keywords": ["string", "..."],',
        '  "combined_role_families": ["string", "..."],',
        '  "notes": "string"',
        "}",
        "",
        "========== PERSONA BLOCKS ==========",
    ]
    for b in blocks:
        lines.append("----------------------------------------------------------------")
        lines.append(f"CODE: {b.get('code')}")
        lines.append(f"LABEL: {b.get('label')}")
        if b.get("triggers"):
            lines.append("CONFIG TRIGGERS: " + "; ".join(str(t) for t in b["triggers"][:25]))
        if b.get("core_stack"):
            lines.append("CONFIG CORE STACK: " + "; ".join(str(s) for s in b["core_stack"]))
        lines.append("RESUME TEXT:")
        lines.append(b.get("resume_text") or "(empty)")
        lines.append("")
    return "\n".join(lines)


def _normalize_resume_keyword_result(raw: dict) -> dict:
    out = dict(raw) if isinstance(raw, dict) else {}
    per = out.get("per_persona")
    if not isinstance(per, list):
        per = []
    cleaned = []
    for row in per:
        if not isinstance(row, dict):
            continue
        code = str(row.get("code") or "").strip().upper()
        if not code:
            continue
        phrases = row.get("linkedin_search_phrases") or row.get("search_phrases") or []
        titles = row.get("target_role_titles") or row.get("roles") or []
        if not isinstance(phrases, list):
            phrases = []
        if not isinstance(titles, list):
            titles = []
        cleaned.append({
            "code": code,
            "label": str(row.get("label") or "").strip(),
            "linkedin_search_phrases": [str(x).strip() for x in phrases if str(x).strip()][:12],
            "target_role_titles": [str(x).strip() for x in titles if str(x).strip()][:12],
            "one_line_focus": str(row.get("one_line_focus") or "").strip(),
        })
    out["per_persona"] = cleaned
    ck = out.get("combined_top_keywords")
    if not isinstance(ck, list):
        ck = []
    out["combined_top_keywords"] = [str(x).strip() for x in ck if str(x).strip()][:20]
    cr = out.get("combined_role_families")
    if not isinstance(cr, list):
        cr = []
    out["combined_role_families"] = [str(x).strip() for x in cr if str(x).strip()][:15]
    out["notes"] = str(out.get("notes") or "").strip()
    out["primary_linkedin_search"] = str(out.get("primary_linkedin_search") or "").strip()[:220]
    return out


def _primary_linkedin_search_pick(norm: dict) -> str:
    """Prefer model's primary string; fallback to combined or per-persona phrases."""
    p = (norm or {}).get("primary_linkedin_search") or ""
    p = str(p).strip()
    if p:
        return p[:220]
    for x in (norm or {}).get("combined_top_keywords") or []:
        s = str(x).strip()
        if s:
            return s[:220]
    for row in (norm or {}).get("per_persona") or []:
        for ph in row.get("linkedin_search_phrases") or []:
            s = str(ph).strip()
            if s:
                return s[:220]
    return ""


def _run_resume_keyword_ai(blocks: list, *, usage_log_endpoint: str):
    """Call provider chain once; returns (normalized_dict, response_text, provider, model_used)."""
    if not blocks:
        raise ValueError("No resume blocks to analyze.")
    prompt = _build_resume_keyword_experiment_prompt(blocks)
    provider = TASK_CHAIN_RESUME_KEYWORDS[0] if TASK_CHAIN_RESUME_KEYWORDS else "groq"
    model_used = ""

    def _on_attempt_error(p, pr, err, lat_ms):
        _log_ai_usage(
            usage_log_endpoint,
            p,
            "unknown",
            pr,
            "",
            lat_ms,
            status="error",
            error_code=str(err),
        )

    started = time.time()
    response_text, model_used, provider, _, _attempted = _run_provider_chain(
        task="resume_keywords",
        chain=TASK_CHAIN_RESUME_KEYWORDS,
        prompt_builder=lambda _p: prompt,
        expect_json=True,
        on_error_log=_on_attempt_error,
    )
    response_text = _strip_markdown_fences(response_text)
    parsed = json.loads(response_text)
    normalized = _normalize_resume_keyword_result(parsed)
    _log_ai_usage(
        usage_log_endpoint,
        provider,
        model_used,
        prompt,
        response_text,
        (time.time() - started) * 1000,
    )
    return normalized, response_text, provider, model_used


# Disable Gemini SDK's internal retry so a 429 (quota) error surfaces immediately
# instead of silently backing off for ~60s and masking itself as a timeout.
_NO_RETRY = gapi_retry.Retry(predicate=lambda _exc: False)


def _generate_with_timeout(model, prompt, timeout_seconds=45):
    # Prevent UI from waiting forever if provider call hangs.
    executor = ThreadPoolExecutor(max_workers=1)
    # Pass retry=None and a request-level timeout so quota / network errors surface fast
    # and don't consume our budget with hidden exponential backoff retries.
    request_options = {
        "retry": _NO_RETRY,
        "timeout": max(5, min(timeout_seconds, 30)),
    }
    future = executor.submit(
        lambda: model.generate_content(prompt, request_options=request_options)
    )
    try:
        return future.result(timeout=timeout_seconds)
    except FuturesTimeoutError as timeout_err:
        future.cancel()
        raise TimeoutError(
            f"AI request timed out after {timeout_seconds} seconds"
        ) from timeout_err
    finally:
        # Do not block request thread waiting for a hung provider call.
        executor.shutdown(wait=False, cancel_futures=True)


def _strip_markdown_fences(text):
    cleaned = (text or "").strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    return cleaned.strip()


_NOTE_GREETING_RE = re.compile(
    r"^(Hi|Hello|Hey|Dear)\s+([A-Za-z][A-Za-z'\-]*)\s*[,\-\u2014]?\s*",
    re.IGNORECASE,
)


_PLACEHOLDER_NAME_RE = re.compile(
    r"^\s*(?:your(?:\s+full|\s+first)?\s+name|your\s+full\s+name|your\s+first\s+name|candidate\s+name)\s*$",
    re.IGNORECASE,
)


def _is_placeholder_name(value):
    text = (value or "").strip()
    if not text:
        return True
    return bool(_PLACEHOLDER_NAME_RE.match(text))


def _safe_candidate_full_name():
    for candidate in (PROFILE.full_name, PROFILE.signoff_name, PROFILE.preferred_name):
        if not _is_placeholder_name(candidate):
            return candidate.strip()
    return "the candidate"


def _safe_signoff_name():
    for candidate in (PROFILE.signoff_name, PROFILE.full_name, PROFILE.preferred_name):
        if not _is_placeholder_name(candidate):
            return candidate.strip()
    return ""


def _scrub_placeholder_identity_text(text):
    cleaned = str(text or "")
    # Convert the most harmful leak pattern into natural language.
    cleaned = re.sub(
        r"\b[Ii]['’]?m\s+Your\s+(?:Full|First)\s+Name\b\s*,?\s*",
        "I'm ",
        cleaned,
    )
    # Remove raw placeholder remnants if they still appear.
    cleaned = re.sub(r"\[?\bYour\s+(?:Full|First)\s+Name\b\]?", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _normalize_connection_note(note, max_chars=300, expected_first_name=None):
    """Clean + length-cap a connection note, defensively fixing two failure
    modes we've seen in real outputs:

    1. Wrong greeting — the model sometimes writes "Hi AJ," (the sender's
       initials) or invents a name when we already know who the recipient
       is. If ``expected_first_name`` is supplied, force the greeting to
       match it.
    2. Mid-word truncation — the previous hard-cut at ``max_chars - 1``
       could leave ugly tails like "...leverage Power B…". Now we prefer
       to stop at the last complete sentence; if no sentence boundary
       fits, we cut at the last space and append an ellipsis.
    """
    cleaned = _scrub_placeholder_identity_text(note)
    cleaned = re.sub(r"\s+", " ", (cleaned or "")).strip()
    if not cleaned:
        return cleaned

    if expected_first_name:
        expected_lower = expected_first_name.strip().lower()
        if expected_lower and expected_lower != "there":
            m = _NOTE_GREETING_RE.match(cleaned)
            if m:
                if m.group(2).lower() != expected_lower:
                    cleaned = f"Hi {expected_first_name}, " + cleaned[m.end():].lstrip()
            else:
                cleaned = f"Hi {expected_first_name}, " + cleaned

    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    if len(cleaned) <= max_chars:
        # If the model left a dangling clause with no terminal punctuation,
        # politely cap it with a period so the note reads as finished.
        if cleaned and cleaned[-1] not in ".!?…":
            cleaned = cleaned.rstrip(" ,;:-") + "."
        return cleaned

    budget = cleaned[:max_chars]
    # Prefer to end at the last complete sentence inside the budget.
    sentence_end = max(budget.rfind(x) for x in (". ", "! ", "? "))
    if sentence_end == -1:
        sentence_end = max(budget.rfind(x) for x in (".", "!", "?"))
    if sentence_end >= int(max_chars * 0.55):
        return budget[: sentence_end + 1].rstrip()

    # No sentence boundary — cut at the last whole word and add an ellipsis.
    last_space = budget.rfind(" ")
    if last_space > 0:
        return budget[:last_space].rstrip(" ,.;:-") + "…"
    return budget.rstrip(" ,.;:-") + "…"


def _as_text(value, fallback=""):
    text = str(value).strip() if value is not None else ""
    return text or fallback


def _as_int(value, fallback=65):
    try:
        return int(round(float(value)))
    except Exception:
        return fallback


def _normalize_status_text(status, bucket="general"):
    """Map whatever the model returned into one of the canonical status
    labels the UI renders with color. Free-tier Llama models frequently
    return semantic phrases like "Acceptable", "Strong alignment", or
    "No concerns" instead of the "GREEN/YELLOW/RED" vocabulary the prompt
    asks for, so this normalizer has a generous keyword list for each
    bucket. Previously, unmatched status values collapsed to "--", which
    is what caused three of five assessment cards to render blank even
    when the details text was fully populated."""
    s = str(status or "").strip().lower()
    if not s:
        return "--"

    # Negation-aware prefilter: phrases like "no concerns", "no rejection",
    # "does not need sponsorship", "no issues", "no blockers" are POSITIVE
    # signals even though they textually contain negative keywords like
    # "concerns" or "rejection". Catch those before the keyword match below
    # so we don't accidentally flip a green signal into yellow/red.
    _positive_negations = (
        "no concern", "no concerns", "no issue", "no issues", "no rejection",
        "no explicit rejection", "no blocker", "no blockers", "no mismatch",
        "no gap", "no gaps", "no red flag", "no red flags", "no obstacle",
        "no obstacles", "no risk", "no risks", "no problem", "no problems",
        "not a blocker", "not a concern", "not an issue",
        "does not need sponsor", "doesn't need sponsor", "no sponsorship needed",
        "no sponsorship required", "sponsorship not required", "opt accepted",
        "opt-accepted", "opt eligible", "stem opt",
    )
    is_positive_negation = any(neg in s for neg in _positive_negations)

    # Signals that consistently indicate each tier, regardless of bucket.
    red_signals = (
        "red", "reject", "deny", "critical", "blocker", "blocked", "disqualif",
        "hard no", "severe", "ineligible", "not a fit", "not a match", "mismatch",
        "dealbreaker",
    )
    yellow_signals = (
        "yellow", "caution", "concern", "concerns", "gap", "partial", "borderline",
        "stretch", "mixed", "moderate", "risk", "unclear", "ambiguous", "limited",
        "somewhat", "some ", "weak", "low-moderate", "watch", "soft", "tbd",
    )
    green_high_signals = (
        "high green", "strong green", "platinum", "excellent", "exceptional",
        "perfect fit", "high match",
    )
    green_signals = (
        "green", "ok", "pass", "eligible", "acceptable", "safe", "aligned",
        "alignment", "matches", "match", "fit", "fits", "positive", "strong",
        "clear", "no concerns", "no concern", "no issues", "meets", "covered",
        "covers", "good", "solid", "suitable", "compatible", "on target",
    )

    # Short-circuit: a positive-negation phrase ("no concerns", "does not
    # need sponsorship", etc.) is unambiguously GREEN regardless of what
    # other keywords also appear in the string.
    if is_positive_negation:
        return "GREEN LIGHT 🟢" if bucket == "sponsorship" else "GREEN 🟢"

    if bucket == "sponsorship":
        # Sponsorship-specific red flags take priority — but only when framed
        # as a rejection, not as "candidate does not need sponsorship".
        if any(k in s for k in ("citizen", "green card", "no opt", "no visa")):
            return "RED LIGHT 🔴"
        if any(k in s for k in red_signals):
            return "RED LIGHT 🔴"
        if any(k in s for k in yellow_signals):
            return "YELLOW LIGHT 🟡"
        if any(k in s for k in green_signals) or any(k in s for k in green_high_signals):
            return "GREEN LIGHT 🟢"
        return "--"

    # General bucket: technical / experience / domain / location.
    if any(k in s for k in red_signals):
        return "RED 🔴"
    if any(k in s for k in yellow_signals):
        return "YELLOW 🟡"
    if any(k in s for k in green_high_signals):
        return "HIGH GREEN 🟢"
    if any(k in s for k in green_signals):
        return "GREEN 🟢"
    return "--"


def _normalize_verdict_line(value, fallback):
    v = str(value or "").strip()
    if not v:
        return fallback
    if "/" in v:
        return v
    return f"{fallback.split('.')[0]}. {v}"


def _normalize_analysis_result(result):
    result = result if isinstance(result, dict) else {}
    verdict = result.get("strategic_verdict")
    if not isinstance(verdict, dict):
        verdict = {}

    # Map common alias keys that some local models return.
    if not result.get("technical_alignment") and isinstance(result.get("technical_fit"), dict):
        result["technical_alignment"] = result.get("technical_fit")
    if not result.get("domain_specialty_gap") and isinstance(result.get("domain_gap"), dict):
        result["domain_specialty_gap"] = result.get("domain_gap")
    if not result.get("experience_seniority") and isinstance(result.get("experience_fit"), dict):
        result["experience_seniority"] = result.get("experience_fit")

    normalized = {
        "company": _as_text(result.get("company"), "Unknown"),
        "position": _as_text(result.get("position"), "Unknown"),
        "resume_code": _as_text(result.get("resume_code"), DEFAULT_RESUME_CODE),
        "resume_rationale": _as_text(result.get("resume_rationale"), ""),
        "sponsorship_legal": result.get("sponsorship_legal") if isinstance(result.get("sponsorship_legal"), dict) else {
            "status": "--",
            "details": "No sponsorship assessment returned.",
        },
        "experience_seniority": result.get("experience_seniority") if isinstance(result.get("experience_seniority"), dict) else {
            "status": "--",
            "details": "No experience assessment returned.",
        },
        "technical_alignment": result.get("technical_alignment") if isinstance(result.get("technical_alignment"), dict) else {
            "status": "--",
            "details": "No technical assessment returned.",
        },
        "domain_specialty_gap": result.get("domain_specialty_gap") if isinstance(result.get("domain_specialty_gap"), dict) else {
            "status": "--",
            "details": "No domain assessment returned.",
        },
        "location_pay": result.get("location_pay") if isinstance(result.get("location_pay"), dict) else {
            "status": "--",
            "details": "No location/pay assessment returned.",
        },
        "fit_score": _as_int(result.get("fit_score"), 65),
        "strategic_verdict": {
            "experience_fit": _as_text(verdict.get("experience_fit"), "5/10. Limited signal; review manually."),
            "technical_fit": _as_text(verdict.get("technical_fit"), "6/10. Some matching skills identified."),
            "growth_fit": _as_text(verdict.get("growth_fit"), "6/10. Potential upside with uncertainty."),
            "recommendation": _as_text(verdict.get("recommendation"), "Apply with review"),
        },
        "suggested_contact": _as_text(result.get("suggested_contact"), "Hiring Manager"),
    }

    normalized["sponsorship_legal"]["status"] = _normalize_status_text(
        normalized["sponsorship_legal"].get("status"), "sponsorship"
    )
    normalized["experience_seniority"]["status"] = _normalize_status_text(
        normalized["experience_seniority"].get("status")
    )
    normalized["technical_alignment"]["status"] = _normalize_status_text(
        normalized["technical_alignment"].get("status")
    )
    normalized["domain_specialty_gap"]["status"] = _normalize_status_text(
        normalized["domain_specialty_gap"].get("status")
    )
    normalized["location_pay"]["status"] = _normalize_status_text(
        normalized["location_pay"].get("status")
    )

    normalized["strategic_verdict"]["experience_fit"] = _normalize_verdict_line(
        normalized["strategic_verdict"].get("experience_fit"),
        "5/10. Limited signal; review manually.",
    )
    normalized["strategic_verdict"]["technical_fit"] = _normalize_verdict_line(
        normalized["strategic_verdict"].get("technical_fit"),
        "6/10. Some matching skills identified.",
    )
    normalized["strategic_verdict"]["growth_fit"] = _normalize_verdict_line(
        normalized["strategic_verdict"].get("growth_fit"),
        "6/10. Potential upside with uncertainty.",
    )

    if normalized["resume_code"] not in RESUMES:
        normalized["resume_code"] = DEFAULT_RESUME_CODE
    normalized["fit_score"] = max(1, min(100, normalized["fit_score"]))

    # Resume rationale: if the model skipped it or gave a vague 1-word answer,
    # synthesize a persona-specific fallback so the UI never shows an empty
    # "Why this resume?" panel. Hard-cap to 280 chars to keep the card tidy.
    rationale = (normalized.get("resume_rationale") or "").strip()
    if len(rationale) < 25:
        rationale = _default_resume_rationale(
            normalized["resume_code"],
            normalized.get("technical_alignment", {}),
            normalized.get("domain_specialty_gap", {}),
        )
    normalized["resume_rationale"] = rationale[:280].rstrip(" ,;:-") or rationale
    return normalized


def _build_rationale_fallbacks_from_personas():
    """Map {code: rationale_fallback} sourced from the persona config."""
    out = {}
    for code in PERSONAS.codes:
        text = PERSONAS.rationale_fallback(code)
        if text:
            out[code] = text
    return out


_RESUME_RATIONALE_FALLBACKS = _build_rationale_fallbacks_from_personas()


_RATIONALE_PLACEHOLDER_DETAILS = {
    "no technical assessment returned.",
    "no domain assessment returned.",
    "no sponsorship assessment returned.",
    "no experience assessment returned.",
    "no location/pay assessment returned.",
    "--",
    "",
}


def _default_resume_rationale(resume_code, technical_alignment, domain_specialty_gap):
    """Build a sensible 'why this resume' sentence when the model didn't return one.

    Prefer a real signal the analysis already surfaced (technical alignment
    details, then domain gap details); fall back to a pure persona pitch when
    no useful signal exists. Skip the placeholder defaults the normalizer
    injects when a field is missing.
    """
    fallback_default = (
        _RESUME_RATIONALE_FALLBACKS.get(DEFAULT_RESUME_CODE)
        or next(iter(_RESUME_RATIONALE_FALLBACKS.values()), "This persona is the closest match to the JD signals.")
    )
    base = _RESUME_RATIONALE_FALLBACKS.get(resume_code, fallback_default)
    for source in (technical_alignment, domain_specialty_gap):
        if not isinstance(source, dict):
            continue
        detail = (source.get("details") or "").strip()
        if detail.lower() in _RATIONALE_PLACEHOLDER_DETAILS:
            continue
        sentence = detail.split(". ")[0].strip().rstrip(".") + "."
        if 20 <= len(sentence) <= 220:
            return f"{base} Signal from this JD: {sentence}"[:280]
    return base[:280]


def _analysis_has_meaningful_content(result):
    if not isinstance(result, dict):
        return False
    fields = [
        result.get("sponsorship_legal"),
        result.get("experience_seniority"),
        result.get("technical_alignment"),
        result.get("domain_specialty_gap"),
        result.get("location_pay"),
    ]
    filled = 0
    for item in fields:
        if isinstance(item, dict):
            details = str(item.get("details", "")).strip()
            status = str(item.get("status", "")).strip()
            # Accept anything with non-empty details (even if status is "--").
            # The previous strict gate triggered avoidable retries on fine answers.
            if details and not details.lower().startswith("no "):
                filled += 1
    # Only retry if the answer is genuinely sparse (fewer than 2 fields filled).
    return filled >= 2


_JD_NOISE_PATTERNS = (
    "equal opportunity",
    "diversity",
    "we celebrate",
    "join our mission",
    "about us",
    "our culture",
    "our values",
    "we offer",
    "perks",
    "401(k)",
    "disability insurance",
    "vision insurance",
    "dental insurance",
    "paid time off",
    "pto",
    "holidays",
    "physical requirements",
    "sedentary work",
    "ada ",
    "benefits found in",
)

_JD_SIGNAL_KEYWORDS = (
    "sponsor", "opt", "ead", "citizen", "green card",
    "years", "experience", "requirements", "responsibilities",
    "must", "required", "qualification", "preferred",
    "salary", "compensation", "$", "/year", "/hr",
    "location", "remote", "hybrid", "onsite", "travel",
    "sql", "python", "excel", "tableau", "power bi", "snowflake",
    "etl", "pipeline", "ml", "machine learning", "statistics",
    "jira", "agile", "scrum", "stakeholder", "report", "dashboard",
)


def _prepare_jd_for_ollama(jd, max_chars=5500):
    text = (jd or "").strip()
    if len(text) <= max_chars:
        return text

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    # Drop obvious boilerplate/noise lines and exact duplicates.
    pruned = []
    seen_lines = set()
    for ln in lines:
        low = ln.lower()
        if any(p in low for p in _JD_NOISE_PATTERNS):
            continue
        if low in seen_lines:
            continue
        seen_lines.add(low)
        pruned.append(ln)

    # Score each remaining line by how many signal keywords it contains.
    scored = []
    for idx, ln in enumerate(pruned):
        low = ln.lower()
        score = sum(1 for k in _JD_SIGNAL_KEYWORDS if k in low)
        # Boost lines that look like requirements/responsibilities headers.
        if any(h in low for h in ("require", "responsib", "qualif", "must have", "preferred", "you will")):
            score += 3
        # Slight bias toward earlier lines (often role title + summary live there).
        if idx < 6:
            score += 1
        scored.append((score, idx, ln))

    scored.sort(key=lambda t: (-t[0], t[1]))

    # Take top-scoring lines until we approach the budget, then sort back into source order.
    chosen = []
    running = 0
    for score, idx, ln in scored:
        if running + len(ln) + 1 > max_chars:
            continue
        chosen.append((idx, ln))
        running += len(ln) + 1
        if running >= max_chars * 0.9:
            break

    chosen.sort(key=lambda t: t[0])
    condensed = "\n".join(ln for _, ln in chosen)
    if not condensed:
        # Fall back to head+tail if scoring produced nothing useful.
        condensed = f"{text[:int(max_chars*0.6)]}\n...\n{text[-int(max_chars*0.3):]}"
    return condensed[:max_chars]


def _build_analyze_prompt(jd, provider="groq"):
    # Cloud LLMs (groq/cerebras/gemini/github) easily handle full JDs in <5s, so we send
    # a generous slice and let them reason over everything. Ollama needs aggressive
    # compression on the light CPU profile to finish within the timeout.
    p = (provider or "").strip().lower()
    if p == "ollama":
        jd_for_prompt = _prepare_jd_for_ollama(jd, max_chars=2400)
    else:
        jd_for_prompt = _prepare_jd_for_ollama(jd, max_chars=8000)
    persona_codes = PERSONAS.code_list_for_prompt()
    priority_line = (
        "PERSONA-SELECTION PRIORITY (when multiple personas match, prefer the FIRST in this list): "
        + " > ".join(PERSONAS.selection_priority)
        if PERSONAS.selection_priority else ""
    )
    return f"""
You are a strict JSON generator for job-fit analysis.
Return ONLY valid JSON with EXACTLY these keys:
company, position, resume_code, resume_rationale, sponsorship_legal, experience_seniority, technical_alignment, domain_specialty_gap, location_pay, fit_score, strategic_verdict, suggested_contact

Requirements:
- resume_code: one of exactly {persona_codes}. You MUST follow the persona triggers in the Resume personas section below. Do NOT pick a persona just because it looks like a safe middle choice. If you cannot justify a pick by naming a specific JD signal that matches a trigger, you are defaulting and must re-evaluate.
- resume_rationale: 1-2 short sentences (max 280 characters) that MUST (a) quote or paraphrase the specific JD signal that matched a trigger, and (b) map it to the chosen persona's lead metric.
- fit_score: integer 1..100
- sponsorship_legal, experience_seniority, technical_alignment, domain_specialty_gap, location_pay: each is an object {{"status": "...", "details": "..."}}.
- STATUS VOCABULARY (CRITICAL — use ONLY these exact strings, nothing else):
    * For sponsorship_legal.status: "GREEN LIGHT" | "YELLOW LIGHT" | "RED LIGHT"
    * For experience_seniority.status, technical_alignment.status, domain_specialty_gap.status, location_pay.status:
        "HIGH GREEN" | "GREEN" | "YELLOW" | "RED"
    * Do NOT invent new status labels like "Acceptable", "Strong alignment", "Safe", "No concerns", "Match", "Positive", "Aligned", etc. — those will be discarded.
    * Use "HIGH GREEN" only when the candidate clearly exceeds the JD requirement. Use "GREEN" for clear alignment. Use "YELLOW" for partial/borderline/stretch. Use "RED" for hard mismatch/disqualification.
- details: one short paragraph explaining WHY you chose that status, quoting a specific JD phrase.
- strategic_verdict has experience_fit, technical_fit, growth_fit, recommendation
- prioritize hard constraints first: sponsorship / work-auth / required experience / mandatory skills
- no markdown, no comments, no extra keys

Example shape (copy this structure exactly, only change the content):
{{
  "sponsorship_legal": {{"status": "GREEN LIGHT", "details": "JD does not require US citizenship; OPT accepted."}},
  "experience_seniority": {{"status": "YELLOW", "details": "JD asks for 3+ years; candidate has ~2 years + MS in progress."}},
  "technical_alignment": {{"status": "GREEN", "details": "SQL + Python + Power BI match the candidate's core stack."}},
  "domain_specialty_gap": {{"status": "GREEN", "details": "No niche domain expertise required."}},
  "location_pay": {{"status": "GREEN", "details": "Hybrid in Charlotte NC; salary range acceptable."}}
}}

{priority_line}

Candidate context:
{USER_CONTEXT}

Resume personas:
{RESUME_CONTEXT}

Job Description:
{jd_for_prompt}
"""


def _ollama_generate(prompt, timeout_seconds=60, expect_json=False, model_name=None, num_predict=None):
    payload = {
        "model": model_name or OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "keep_alive": OLLAMA_KEEP_ALIVE,
        "options": {
            "num_ctx": OLLAMA_NUM_CTX,
            "num_thread": OLLAMA_NUM_THREAD,
            "num_gpu": OLLAMA_NUM_GPU,
            "temperature": OLLAMA_TEMPERATURE,
            "num_predict": num_predict or OLLAMA_NUM_PREDICT_DEFAULT,
        },
    }
    if expect_json:
        payload["format"] = "json"

    req = url_request.Request(
        f"{OLLAMA_URL}/api/generate",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with url_request.urlopen(req, timeout=timeout_seconds) as resp:
            body = resp.read().decode("utf-8")
    except url_error.HTTPError as e:
        details = e.read().decode("utf-8", errors="replace") if hasattr(e, "read") else str(e)
        raise Exception(f"Ollama HTTP {e.code}: {details}") from e
    except url_error.URLError as e:
        raise Exception(f"Ollama unreachable/not running: {e}") from e
    except TimeoutError as e:
        raise TimeoutError(f"AI request timed out after {timeout_seconds} seconds") from e

    parsed = json.loads(body)
    response_text = parsed.get("response", "")
    if not response_text:
        raise Exception("Ollama returned an empty response")
    return response_text


def _ai_generate_text(prompt, timeout_seconds=45, expect_json=False, task="general"):
    return _ai_generate_text_with_provider(prompt, AI_PROVIDER, timeout_seconds, expect_json, task)[0]


def _ai_generate_text_with_provider(prompt, provider, timeout_seconds=45, expect_json=False, task="general"):
    provider = (provider or "").strip().lower()
    if provider == "groq":
        if not GROQ_API_KEY:
            raise Exception("Groq API key not configured (set GROQ_API_KEY in .env)")
        model_name = _model_for_groq(task)
        text = _call_openai_compatible(
            GROQ_BASE_URL, GROQ_API_KEY, model_name, prompt,
            timeout_seconds=timeout_seconds, expect_json=expect_json,
            temperature=_temperature_for_task(task),
            max_tokens=1800 if task in ("analyze", "cover", "draft") else 700,
        )
        return text, model_name
    if provider == "cerebras":
        if not CEREBRAS_API_KEY:
            raise Exception("Cerebras API key not configured (set CEREBRAS_API_KEY in .env)")
        model_name = _model_for_cerebras(task)
        text = _call_openai_compatible(
            CEREBRAS_BASE_URL, CEREBRAS_API_KEY, model_name, prompt,
            timeout_seconds=timeout_seconds, expect_json=expect_json,
            temperature=_temperature_for_task(task),
            max_tokens=1800 if task in ("analyze", "cover", "draft") else 700,
        )
        return text, model_name
    if provider == "github":
        if not GITHUB_MODELS_TOKEN:
            raise Exception("GitHub Models token not configured (set GITHUB_MODELS_TOKEN in .env)")
        model_name = _model_for_github(task)
        text = _call_openai_compatible(
            GITHUB_MODELS_BASE_URL, GITHUB_MODELS_TOKEN, model_name, prompt,
            timeout_seconds=timeout_seconds, expect_json=expect_json,
            temperature=_temperature_for_task(task),
            max_tokens=1800 if task in ("analyze", "cover", "draft") else 700,
        )
        return text, model_name
    if provider == "gemini":
        generation_config = {"response_mime_type": "application/json"} if expect_json else None
        model_list = [GENAI_MODEL] + [m for m in GEMINI_MODEL_CANDIDATES if m != GENAI_MODEL]
        last_error = None
        for model_name in model_list:
            try:
                model = genai.GenerativeModel(model_name, generation_config=generation_config)
                text = _generate_with_timeout(model, prompt, timeout_seconds=timeout_seconds).text.strip()
                return text, model_name
            except Exception as e:
                last_error = e
                err_str = str(e).upper()
                # Only cycle to the next Gemini candidate on "model unavailable"-style failures.
                # For timeouts, quota (429), or network errors, cycling wastes time — bail out
                # so the caller can fall back to the OTHER provider (Ollama) immediately.
                recoverable = (
                    "NOT FOUND" in err_str
                    or "404" in err_str
                    or "PERMISSION_DENIED" in err_str
                    or "DEPRECATED" in err_str
                )
                if not recoverable:
                    break
        raise last_error or Exception("No Gemini model succeeded")
    if provider == "ollama":
        model_name = OLLAMA_MODEL_ANALYZE if task == "analyze" else OLLAMA_MODEL
        num_predict = OLLAMA_NUM_PREDICT_ANALYZE if task == "analyze" else OLLAMA_NUM_PREDICT_DEFAULT
        text = _ollama_generate(
            prompt,
            timeout_seconds=timeout_seconds,
            expect_json=expect_json,
            model_name=model_name,
            num_predict=num_predict,
        ).strip()
        return text, model_name
    raise Exception(f"Unsupported provider: {provider}")


_PRONOUN_TOKENS = {
    "she", "her", "hers", "he", "him", "his", "they", "them", "their", "theirs",
    "ze", "zir", "zem", "ey", "em", "eir",
}
_CONNECTION_NOISE_TOKENS = {
    "1st", "2nd", "3rd", "1st degree", "2nd degree", "3rd degree", "connection",
    "follower", "followers", "following", "mutual", "mutuals",
    "dr", "mr", "mrs", "ms", "miss", "prof", "professor",  # titles handled below
}


def _extract_clean_contact_name(contact_name):
    """Pull a clean 'First [Middle] Last' name out of messy LinkedIn pastes.

    Handles inputs like:
        "Jane Doe She/Her · 3rd · Recruiter @ Acme Corp..."
        "John Smith\n1st degree connection · 1st\nProduct Analytics..."
        "Alex Lee (She/Her) · 1st"
    """
    raw_contact_input = (contact_name or "").strip()
    if not raw_contact_input:
        return "", ""

    # 1) Take the first non-empty line — most LinkedIn pastes put the name there.
    first_non_empty_line = next(
        (line.strip() for line in raw_contact_input.splitlines() if line.strip()),
        raw_contact_input,
    )

    # 2) Strip everything after structural separators that always follow the name:
    #    "·" (LinkedIn middle dot), "|", ",", "@", "-", "—", "–", "•"
    candidate = re.split(r"[\u00B7\|@,\-\u2014\u2013\u2022]", first_non_empty_line, maxsplit=1)[0]

    # 3) Strip parenthetical pronouns: "Agness Daszykowski (She/Her)" -> "Agness Daszykowski"
    candidate = re.sub(r"\s*\([^)]*\)\s*", " ", candidate)

    # 4) Tokenize, then drop pronouns / connection-degree noise / common titles.
    raw_tokens = re.findall(r"[A-Za-z][A-Za-z'\-]*", candidate)
    keep = []
    for tok in raw_tokens:
        low = tok.lower().rstrip(".")
        if low in _PRONOUN_TOKENS:
            continue
        if low in _CONNECTION_NOISE_TOKENS:
            continue
        # If we hit a token that already ends in /something (LinkedIn pronoun glob)
        # treat the whole thing as a stop; only the prefix before / is real.
        keep.append(tok)
        if len(keep) >= 3:  # most names are 1-3 tokens
            break

    clean_contact_name = " ".join(keep).strip()
    if not clean_contact_name:
        # Last-ditch fallback: take first 1-2 alphabetic words from the raw input.
        fallback_tokens = re.findall(r"[A-Za-z][A-Za-z'\-]*", raw_contact_input)
        clean_contact_name = " ".join(fallback_tokens[:2]).strip() or raw_contact_input
    return raw_contact_input, clean_contact_name


def _build_connection_note_prompt(company, position, contact_name, resume_code, jd):
    raw_contact_input, clean_contact_name = _extract_clean_contact_name(contact_name)
    first_name = clean_contact_name.split(" ")[0] if clean_contact_name else "there"
    note_greeting = f"Hi {first_name},"

    candidate_full_name = _safe_candidate_full_name()
    note_char_limit = APP_SETTINGS.limit("connection_note_max_chars", 300)

    affinity_label = (PROFILE.networking.get("alumni_button_label") or "").strip()
    affinity_keywords = []
    if PROFILE.alumni_school_slug:
        for tok in re.split(r"[-_\s]+", PROFILE.alumni_school_slug.lower()):
            if tok and len(tok) > 2:
                affinity_keywords.append(tok)
    if affinity_label:
        for tok in re.split(r"\s+", affinity_label.lower()):
            if tok and len(tok) > 2:
                affinity_keywords.append(tok)
    networking_signals = f"{raw_contact_input}\n{jd}".lower()
    has_affinity = (
        PROFILE.alumni_search_enabled
        and any(k in networking_signals for k in affinity_keywords)
    )
    if has_affinity and affinity_label:
        networking_instruction = (
            f"You MAY reference a {affinity_label} angle ONLY if alignment is explicit in the contact context or JD."
        )
    else:
        networking_instruction = (
            "Do NOT imply alumni / affinity-network overlap unless explicitly supported by the provided context."
        )

    banned_phrases_block = ""
    if PROFILE.banned_phrases:
        banned_phrases_block = "- Do NOT use these banned phrases: " + ", ".join(
            f'"{p}"' for p in PROFILE.banned_phrases
        )

    prompt = f"""
You are drafting a LinkedIn connection request note for {candidate_full_name}.

Recipient first name (use ONLY this in the greeting): "{first_name}"
Recipient full name (reference only): "{clean_contact_name}"
Raw recipient context for inferring their role/team (NEVER quote verbatim): "{raw_contact_input}"
Target role: "{position}" at "{company}"

CANDIDATE BIO + RESUME PERSONA "{resume_code}" (use these exact facts; do NOT invent new ones):
{RESUME_CONTEXT}

VOICE: Match the candidate's actual seniority and professional voice as conveyed by the bio above. Do NOT inflate to senior/executive tone if the bio describes an early-career or student profile. The "persona voice" labels in the resume personas section are INTERNAL STYLE HINTS — do NOT recite them verbatim in the note. Use plain self-descriptions such as "I'm a [role/program] focused on...", "I specialize in...".

PATTERN TO FOLLOW — every great note hits these beats in order:
  1) Greeting: "{note_greeting}"
  2) ONE-LINE self-positioning in the candidate's authentic voice. Grounded, NOT grandiose.
  3) ONE specific reference to {company} or the team implied by the JD (NOT generic phrases like "your company" or "data analysis"). Pull a real signal from the JD: team name, product, mission line, function, or specific tech.
  4) ONE low-pressure reason to connect (e.g. "would love to connect to see how you...", "would love to stay connected").

HARD RULES:
- HARD CHARACTER LIMIT: {note_char_limit} characters total including spaces. Target {max(60, note_char_limit - 70)}-{max(80, note_char_limit - 15)} so you always have room to end on a complete sentence.
- ENDING: The note MUST end with a complete sentence and terminal punctuation (., !, or ?). NEVER end mid-word, mid-phrase, or with an ellipsis.
- Start EXACTLY with "{note_greeting}" verbatim — copy the first name character-for-character. Do NOT substitute initials, nicknames, last names, or any other name.
- DO NOT recite the internal persona-voice labels word-for-word. Describe WHAT the candidate does, do not paste a label like a job title.
- Honor any "exact_phrasing_required" notes in the candidate's education list — never substitute a different degree title or institution.
- AI-TOOLS RULE: If the JD names any AI tool (LLMs, agents, prompt engineering, Cursor, ChatGPT, Gemini, Copilot, etc.) AND the candidate's bio includes AI-tools experience, work in ONE specific AI fact from the bio. Do NOT fabricate AI experience the bio does not list.
{banned_phrases_block}
- Do NOT use hashtags, emojis (unless one fits naturally — default to none), bullets, or line breaks.
- Do NOT name-drop fabricated tools/projects/numbers. Stick to the bio facts above.
- {networking_instruction}

Return PURE JSON with exactly one key: "connection_note". No markdown fences.

Job Description:
{jd}
"""
    return prompt


@app.route("/api/ask", methods=["POST"])
def ask_tracker_agent():
    if not ENABLE_FOLLOWUP_AGENT:
        return jsonify({"error": "Follow-up assistant is disabled."}), 403

    data = request.json or {}
    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({"error": "Question cannot be empty."}), 400

    rows = _load_recent_tracker_rows(limit=60)
    tracker_context = _build_tracker_context(rows)

    provider = TASK_CHAIN_ASK[0] if TASK_CHAIN_ASK else "groq"
    model_used = ""
    started = time.time()
    try:
        prompt = f"""
You are a follow-up assistant for a job application tracker.
Answer only from the tracker data below.
If the tracker does not contain the requested fact, say clearly: 'I cannot find that in your tracker yet.'
Be concise and factual.

Tracker data:
{tracker_context}

User question:
{question}
"""
        def _on_attempt_error(p, pr, err, lat_ms):
            _log_ai_usage("/api/ask", p, "unknown", pr, "", lat_ms, status="error", error_code=str(err))

        answer, model_used, provider, _, _ = _run_provider_chain(
            task="ask",
            chain=TASK_CHAIN_ASK,
            prompt_builder=lambda _p: prompt,
            expect_json=False,
            on_error_log=_on_attempt_error,
        )
        _log_ai_usage("/api/ask", provider, model_used, prompt, answer, (time.time() - started) * 1000)
        return jsonify({"answer": answer, "rows_considered": len(rows)})
    except Exception as e:
        _log_ai_usage(
            "/api/ask",
            provider,
            model_used or ("unknown"),
            prompt if 'prompt' in locals() else "",
            "",
            (time.time() - started) * 1000,
            status="error",
            error_code=str(e),
        )
        msg, code = _friendly_ai_error(e, "Failed to answer follow-up question")
        return jsonify({"error": msg}), code


@app.route("/api/analyze", methods=["POST"])
def analyze_fit():
    data = request.json
    jd = data.get("jd", "")
    
    if not jd.strip():
        return jsonify({"error": "Job Description cannot be empty"}), 400

    summary = _read_usage_summary_today()
    if summary["analyze_today"] >= DAILY_ANALYZE_HARD_CAP:
        return jsonify({
            "error": f"Daily analyze hard cap reached ({DAILY_ANALYZE_HARD_CAP}). Try again tomorrow or switch provider policy."
        }), 429

    provider = TASK_CHAIN_ANALYZE[0] if TASK_CHAIN_ANALYZE else "groq"
    model_used = ""
    prompt = ""
    started = time.time()
    try:
        def _on_attempt_error(p, pr, err, lat_ms):
            _log_ai_usage("/api/analyze", p, "unknown", pr, "", lat_ms, status="error", error_code=str(err))

        response_text, model_used, provider, prompt, attempted = _run_provider_chain(
            task="analyze",
            chain=TASK_CHAIN_ANALYZE,
            prompt_builder=lambda p: _build_analyze_prompt(jd, provider=p),
            expect_json=True,
            on_error_log=_on_attempt_error,
        )
        response_text = _strip_markdown_fences(response_text)
        
        try:
            result = json.loads(response_text)
        except Exception as e:
            print("Failed JSON:", response_text)
            raise Exception(f"AI produced invalid JSON layout: str({e})")
        
        result = _normalize_analysis_result(result)
        if provider == "ollama" and not _analysis_has_meaningful_content(result):
            retry_prompt = _build_analyze_prompt(jd, provider="ollama")
            # Use the dedicated Ollama analyze timeout for the retry too — using the much
            # shorter Gemini timeout here was causing chained timeouts (75s + 30s = 105s).
            retry_text, _ = _ai_generate_text_with_provider(
                retry_prompt, "ollama", timeout_seconds=AI_TIMEOUT_ANALYZE_OLLAMA,
                expect_json=True, task="analyze",
            )
            retry_text = _strip_markdown_fences(retry_text)
            try:
                retry_result = json.loads(retry_text)
                retry_result = _normalize_analysis_result(retry_result)
                if _analysis_has_meaningful_content(retry_result):
                    result = retry_result
            except Exception:
                pass

        # Enforce honest fit_score caps when the model returns contradictory numbers
        fit = result.get("fit_score")
        if isinstance(fit, (int, float)):
            fit = int(round(fit))
            exp = result.get("experience_seniority") or {}
            dom = result.get("domain_specialty_gap") or {}
            es = (str(exp.get("status", ""))).upper()
            ds = (str(dom.get("status", ""))).upper()
            if "RED" in es:
                fit = min(fit, 65)
            elif "YELLOW" in es:
                fit = min(fit, 78)
            if "RED" in ds:
                fit = min(fit, 72)
            elif "YELLOW" in ds and "RED" not in es:
                fit = min(fit, 82)
            result["fit_score"] = max(1, min(100, fit))

        # Log every successful analysis so the tracker has a row even without outreach/email
        try:
            company = (result.get("company") or "Unknown").strip() or "Unknown"
            position = (result.get("position") or "Unknown").strip() or "Unknown"
            rc = result.get("resume_code") or DEFAULT_RESUME_CODE
            fs = result.get("fit_score")
            fs_part = f"fit {fs}%" if isinstance(fs, (int, float)) else "fit n/a"
            log_job(
                company=company,
                position=position,
                resume_used=rc,
                contact_name="N/A",
                role="N/A",
                platform="Career Command Center",
                message=f"JD analyzed — {fs_part}",
                status="Analyzed",
                tracker_path=JOB_TRACKER_PATH,
            )
        except Exception as log_err:
            print("Analysis CSV log skipped:", log_err)
            
        latency_ms = int((time.time() - started) * 1000)
        _log_ai_usage("/api/analyze", provider, model_used, prompt, response_text, latency_ms)
        # Expose provider telemetry to the UI so the user can see which model
        # actually won the provider chain (helpful when debugging "why does
        # this analysis look off?" or "which AI did this?").
        result["_meta"] = {
            "provider": provider,
            "model": model_used or "unknown",
            "latency_ms": latency_ms,
            "chain_attempted": attempted or [provider],
        }
        return jsonify(result)

    except Exception as e:
        _log_ai_usage(
            "/api/analyze",
            provider,
            model_used or "unknown",
            prompt if 'prompt' in locals() else "",
            "",
            (time.time() - started) * 1000,
            status="error",
            error_code=str(e),
        )
        msg, code = _friendly_ai_error(e, "Failed to parse analysis")
        return jsonify({"error": msg}), code


@app.route("/api/draft", methods=["POST"])
def draft_email():
    data = request.json
    jd = data.get("jd", "")
    company = data.get("company", "the company")
    position = data.get("position", "the open position")
    contact_name = data.get("contact_name", "")
    resume_code = data.get("resume_code", DEFAULT_RESUME_CODE)

    if not jd.strip():
        return jsonify({"error": "Job Description missing"}), 400
    if not contact_name.strip():
        return jsonify({"error": "Enter the target person's name first so the connection note can be personalized."}), 400

    provider = TASK_CHAIN_DRAFT[0] if TASK_CHAIN_DRAFT else "groq"
    model_used = ""
    started = time.time()
    try:
        raw_contact_input, clean_contact_name = _extract_clean_contact_name(contact_name)

        candidate_full_name = _safe_candidate_full_name()
        signoff_name = _safe_signoff_name() or (
            candidate_full_name if candidate_full_name != "the candidate" else ""
        )
        note_char_limit = APP_SETTINGS.limit("connection_note_max_chars", 300)
        body_min = APP_SETTINGS.limit("email_body_min_words", 140)
        body_max = APP_SETTINGS.limit("email_body_max_words", 230)

        greeting = (
            f"Dear {clean_contact_name},"
            if clean_contact_name and clean_contact_name.lower() != "n/a"
            else f"Hi {company} Talent Team,"
        )
        firstName = clean_contact_name.split(" ")[0] if clean_contact_name else "Team"
        note_greeting = f"Hi {firstName},"

        # Profile-driven affinity / networking instruction (no hardcoded school).
        affinity_label = (PROFILE.networking.get("alumni_button_label") or "").strip()
        affinity_keywords = []
        if PROFILE.alumni_school_slug:
            for tok in re.split(r"[-_\s]+", PROFILE.alumni_school_slug.lower()):
                if tok and len(tok) > 2:
                    affinity_keywords.append(tok)
        if affinity_label:
            for tok in re.split(r"\s+", affinity_label.lower()):
                if tok and len(tok) > 2:
                    affinity_keywords.append(tok)
        networking_signals = f"{raw_contact_input}\n{jd}".lower()
        has_affinity = (
            PROFILE.alumni_search_enabled
            and any(k in networking_signals for k in affinity_keywords)
        )
        if has_affinity and affinity_label:
            networking_instruction = (
                f"You may reference a {affinity_label} angle ONLY if the recipient is explicitly affiliated in the contact context or JD."
            )
        else:
            networking_instruction = (
                "Do NOT imply alumni / affinity-network overlap unless explicitly supported by the provided contact context or JD."
            )

        banned_phrases_block = ""
        if PROFILE.banned_phrases:
            banned_phrases_block = "- NEVER use these banned phrases: " + ", ".join(
                f'"{p}"' for p in PROFILE.banned_phrases
            )

        draft_prompt = f"""
You are drafting cold outreach for {candidate_full_name} applying for "{position}" at "{company}".
You will produce three artifacts: a LinkedIn connection note, an email subject, and an email body.

Recipient first name (use ONLY this in greetings): "{firstName}"
Recipient full name (use exactly this in the email "Dear …" greeting): "{clean_contact_name}"
Raw recipient context (infer their team/seniority — never quote verbatim): "{raw_contact_input}"

Email greeting MUST be exactly: "{greeting}"
Connection note greeting MUST be exactly: "{note_greeting}"

CANDIDATE BIO + RESUME PERSONA "{resume_code}" (these are the ONLY facts you may cite — do NOT invent new numbers, employers, or tools):
{RESUME_CONTEXT}

{networking_instruction}

================ HARD RULES ================

CONNECTION NOTE ("connection_note"):
- HARD LIMIT: {note_char_limit} characters total. Target {max(60, note_char_limit - 60)}-{max(80, note_char_limit - 10)}.
- Start EXACTLY with "{note_greeting}".
- 4 beats in order: greeting -> one-line self-positioning in the persona-{resume_code} voice -> ONE specific reference to {company} (team/product/mission inferred from the JD) -> ONE low-pressure reason to connect.
- No hashtags, no bullets, no emojis (unless one truly fits), no line breaks.
- Do NOT mention GPA in the note (save numbers for the email).

EMAIL SUBJECT ("subject"):
- Specific and outcome-led. Reference the role + a real differentiator from the bio.
- AVOID generic patterns like "Application for X" or "Data Analysis Expertise for Y".
- 6-12 words ideal.

EMAIL BODY ("body"):
- {body_min}-{body_max} words.
- Greeting line = exactly "{greeting}" on its own line, followed by a blank line.
- Open with WHY this role + which specific team/function from the JD (use a real signal pulled from the JD — team name, product, mission line, etc.).
- Middle paragraph MUST cite at least TWO concrete proof points from the bio above (only metrics, employers, or tools that actually appear in the bio). Tie each proof point to an explicit JD requirement.
- One short paragraph that is either (a) a low-pressure ask for the recipient's perspective on the team OR (b) a confident statement of fit.
- Sign-off: "{signoff_name}" on its own line. (You may add "Best regards," on the line above if the email feels relationship-warm.)
- Mention that the resume is attached.
- NEVER write placeholders like [Your Phone Number], [Date], [Address].
- NEVER paste recipient pronouns / connection-degree / title into the greeting line.

HARD FACTS (apply to BOTH the note and the email body):
- Honor every "exact_phrasing_required" note in the candidate's education list — never substitute a different degree title or institution.
- Cite ONLY metrics, employers, schools, and tools that appear in the CANDIDATE BIO above.

AI-TOOLS RULE (conditional — scan the JD first):
- If the JD mentions any AI tool / LLM / agentic workflow signals AND the candidate's bio includes AI-tools experience, work in ONE concrete AI-tools fact from the bio. Otherwise, do NOT mention AI tools.
{banned_phrases_block}

OUTPUT FORMAT:
Return PURE JSON with EXACTLY these 3 keys: "subject", "body", "connection_note".
Use single quotes inside string values — never inner double quotes. No markdown fences.

Job Description:
{jd}
"""
        def _on_attempt_error(p, pr, err, lat_ms):
            _log_ai_usage("/api/draft", p, "unknown", pr, "", lat_ms, status="error", error_code=str(err))

        response_text, model_used, provider, _, _ = _run_provider_chain(
            task="draft",
            chain=TASK_CHAIN_DRAFT,
            prompt_builder=lambda _p: draft_prompt,
            expect_json=True,
            on_error_log=_on_attempt_error,
        )
        response_text = _strip_markdown_fences(response_text)
        
        try:
            result = json.loads(response_text)
        except Exception as e:
            print("Failed Draft JSON:", response_text)
            raise Exception(f"AI produced invalid JSON layout: str({e})")

        # Guardrail for UI/platform limit.
        if result.get("connection_note"):
            result["connection_note"] = _enforce_hard_facts(_normalize_connection_note(
                result.get("connection_note"),
                expected_first_name=firstName if firstName and firstName.lower() != "team" else None,
            ))
        # Degree-accuracy safety net for the email body/subject too.
        if result.get("body"):
            result["body"] = _enforce_hard_facts(_scrub_placeholder_identity_text(result["body"]))
        if result.get("subject"):
            result["subject"] = _enforce_hard_facts(_scrub_placeholder_identity_text(result["subject"]))

        _log_ai_usage("/api/draft", provider, model_used, draft_prompt, response_text, (time.time() - started) * 1000)
        return jsonify(result)

    except Exception as e:
        _log_ai_usage(
            "/api/draft",
            provider,
            model_used or "unknown",
            draft_prompt if 'draft_prompt' in locals() else "",
            "",
            (time.time() - started) * 1000,
            status="error",
            error_code=str(e),
        )
        msg, code = _friendly_ai_error(e, "Failed to draft outreach")
        return jsonify({"error": msg}), code


@app.route("/api/connection-note", methods=["POST"])
def draft_connection_note():
    data = request.json
    jd = data.get("jd", "")
    company = data.get("company", "the company")
    position = data.get("position", "the open position")
    contact_name = data.get("contact_name", "")
    resume_code = data.get("resume_code", DEFAULT_RESUME_CODE)

    if not jd.strip():
        return jsonify({"error": "Job Description missing"}), 400
    if not contact_name.strip():
        return jsonify({"error": "Enter the target person's name first so the connection note can be personalized."}), 400

    prompt = _build_connection_note_prompt(company, position, contact_name, resume_code, jd)
    started = time.time()

    def _on_attempt_error(p, pr, err, lat_ms):
        _log_ai_usage("/api/connection-note", p, "unknown", pr, "", lat_ms, status="error", error_code=str(err))

    try:
        response_text, model_used, provider, _, _ = _run_provider_chain(
            task="note",
            chain=TASK_CHAIN_NOTE,
            prompt_builder=lambda _p: prompt,
            expect_json=True,
            on_error_log=_on_attempt_error,
        )
        response_text = _strip_markdown_fences(response_text)
        try:
            result = json.loads(response_text)
        except Exception:
            result = {"connection_note": response_text.strip()}
        note = (result.get("connection_note") or "").strip()
        if not note:
            raise Exception("AI did not return a connection note")
        _log_ai_usage(
            "/api/connection-note", provider, model_used, prompt, response_text,
            (time.time() - started) * 1000,
        )
        _, clean_contact = _extract_clean_contact_name(contact_name)
        expected_first = clean_contact.split(" ")[0] if clean_contact else "there"
        return jsonify({
            "connection_note": _enforce_hard_facts(
                _normalize_connection_note(note, expected_first_name=expected_first)
            ),
            "provider_used": provider,
        })
    except Exception as last_err:
        emergency_note = _build_emergency_connection_note(company, position, contact_name)
        return jsonify({
            "connection_note": emergency_note,
            "provider_used": "template_fallback",
            "warning": f"AI providers unavailable: {last_err}",
        })


_LEGAL_SUFFIX_RE = re.compile(
    r"[,\s]+(?:L\.?L\.?C\.?|Inc\.?|Incorporated|Corp\.?|Corporation|Ltd\.?|Limited|"
    r"L\.?L\.?P\.?|PLC|GmbH|S\.?A\.?|S\.?p\.?A\.?|Pty\s+Ltd|Pte\s+Ltd|N\.?V\.?|B\.?V\.?|AG)\s*$",
    flags=re.IGNORECASE,
)


def _shorten_company_for_copy(name):
    """'Acme Asset Management, LLC' -> 'Acme Asset Management'.

    Used so the cover letter doesn't repeat a legal suffix like 'LLC' three
    times in the same page. The full form is still used once on the
    greeting or first mention.
    """
    if not name:
        return ""
    trimmed = name.strip()
    if not trimmed:
        return ""
    cleaned = _LEGAL_SUFFIX_RE.sub("", trimmed).strip().rstrip(",").strip()
    return cleaned or trimmed


def _company_first_word(name):
    """'Acme Electric Company' -> 'Acme'. Used for very late mentions where
    an even shorter alias reads more naturally."""
    short = _shorten_company_for_copy(name)
    tokens = re.findall(r"[A-Za-z0-9]+", short)
    return tokens[0] if tokens else short


# Public URLs the candidate wants appended at the bottom of every cover letter.
# Sourced from the configured profile so each user's contact links are theirs.
COVER_LETTER_CONTACT_URLS = list(PROFILE.contact_urls_for_cover_letter)


def _append_contact_urls_to_letter(letter):
    """Append portfolio + LinkedIn URLs after the sign-off, once. Idempotent —
    won't duplicate URLs if the AI already included them somewhere."""
    text = (letter or "").rstrip()
    if not text:
        return text
    if not COVER_LETTER_CONTACT_URLS:
        return text
    lower = text.lower()
    urls_to_add = [u for u in COVER_LETTER_CONTACT_URLS if u.lower() not in lower]
    if not urls_to_add:
        return text
    return text + "\n\n" + "\n".join(urls_to_add)


def _build_cover_letter_prompt(company, position, resume_code, jd, tone="professional"):
    jd_for_prompt = _prepare_jd_for_ollama(jd, max_chars=4500)
    tone_hint = {
        "professional": "warm, confident, professional",
        "enthusiastic": "energetic, mission-driven, passionate but grounded",
        "concise": "tight, senior, outcome-first with zero filler",
    }.get((tone or "professional").lower(), "warm, confident, professional")

    company_full = (company or "").strip() or "the company"
    company_short = _shorten_company_for_copy(company_full) or company_full
    company_first = _company_first_word(company_full) or company_short

    candidate_full_name = _safe_candidate_full_name()
    signoff_name = _safe_signoff_name() or (candidate_full_name if candidate_full_name != "the candidate" else "")
    cover_min = APP_SETTINGS.limit("cover_letter_min_words", 320)
    cover_max = APP_SETTINGS.limit("cover_letter_max_words", 450)

    banned_phrases_clause = ""
    if PROFILE.banned_phrases:
        banned_phrases_clause = " Do NOT use these banned phrases: " + ", ".join(
            f'"{p}"' for p in PROFILE.banned_phrases
        )

    return f"""
You are writing a tailored cover letter for {candidate_full_name} applying for "{position}" at "{company_full}".

COMPANY NAMES (use these EXACT forms — do not paste any other variant):
  - Full legal name (use AT MOST ONCE, in the opening paragraph's first mention): "{company_full}"
  - Short form (use for EVERY OTHER mention, including the greeting line and the bullet-section header): "{company_short}"
  - First-word alias (OK to use one or two times for late possessive mentions like "{company_first}'s mission"): "{company_first}"
The legal suffix (LLC, Inc, Corp, Ltd, LLP, etc.) must appear AT MOST ONCE in the entire letter.

================ HARD RULES (failure to follow = bad output) ================

OUTPUT FORMAT:
- Return PURE JSON with EXACTLY one key: "cover_letter".
- The value is plain text. No markdown headers, no code fences.
- Use REAL line breaks (\\n) inside the JSON string. Each paragraph MUST be separated from the next by a blank line (two consecutive \\n). Bullet list lines start with "- " (hyphen + space).

STRUCTURE (MANDATORY — these 6 blocks, in order, with blank lines between every block):
  1) Greeting line: "Dear {company_short} Hiring Team," — or "To the {company_short} Team," if more natural. NEVER include the legal suffix here.
  2) Opening paragraph (3-5 sentences): Position {candidate_full_name} (use the candidate's actual education / background as written in the bio — never substitute a different degree title), state which specific role at {company_full} you are writing about (this is the ONE place where the full legal form may appear), and reference one concrete signal from the JD (team name, location, mission line).
  3) Evidence paragraph (3-5 sentences): Cite 2-3 specific quantified wins from the bio and map each to an explicit JD responsibility or requirement. Use the persona-{resume_code} framing implied by the resume personas section. Refer to the company as "{company_short}" or "{company_first}" here.
  4) Strategic-storytelling paragraph (3-5 sentences): Reference one concrete project from the bio that shows stakeholder + cross-functional translation. Tie it to a relevant element of the JD.
  5) Bullet section titled exactly: "Why I am the right fit for {company_short}:". EXACTLY 3 bullet items, each on its own line, each starting with "- " (hyphen + space) followed by a short Label + colon + 1-sentence explanation. Example: "- Technical Rigor: I bring advanced proficiency in SQL, Python, and Power BI."
  6) Closing paragraph (1-3 sentences) followed by a blank line, then a sign-off line: "{signoff_name}" on its own line.

NO MARKDOWN: Do NOT use any markdown formatting anywhere. No asterisks for bold (**like this**), no underscores for italic, no hashes for headings (#), no backticks. The output is plain prose that will be rendered into a PDF directly.

VOICE: {tone_hint}.{banned_phrases_clause}

LENGTH: {cover_min}-{cover_max} words total. Hard ceiling {cover_max + 30} words.

FACTS: Only cite numbers, employers, schools, and tools that appear in the CANDIDATE BIO below. Do NOT fabricate company names, project names, or metrics. Honor any "exact_phrasing_required" notes in the candidate's education list — never substitute a different degree title or institution.

AI-TOOLS RULE (conditional — check the JD before writing):
- If the JD mentions AI tools / LLMs / agents / prompt engineering / agentic workflows AND the candidate's bio includes AI-tools experience, work in ONE concrete AI-tools fact from the bio.
- Otherwise, do NOT mention AI tools — keep the letter focused on the JD's actual priorities.

NO PLACEHOLDERS: Never write [Your Address], [Date], [Phone], [Hiring Manager Name]. Never include a letterhead — the document body starts at the greeting. Do NOT include a "Portfolio:" or "LinkedIn:" line — the candidate's URLs are appended automatically after your sign-off by the server.

SPONSORSHIP: Use the candidate's bio to assess work-authorization framing. Do not over-discuss sponsorship unless the JD explicitly raises a hard requirement.

================ INPUTS ================

Candidate context:
{USER_CONTEXT}

Resume persona to emphasize ({resume_code}):
{RESUME_CONTEXT}

Job Description:
{jd_for_prompt}
"""


_BULLET_LINE_RE = re.compile(r"^\s*[\u2022\u25CF\u25AA\-\*]\s+")


_HARD_FACT_PATCHES = []  # Populated below from profile.json — see _build_hard_fact_patches.


def _build_hard_fact_patches(profile):
    """Build regex-based fact-correction patches from the candidate profile.

    For each education entry that defines an "exact_phrasing_required", we
    construct a defensive guard so a provider that slips into a wrong degree
    title (e.g. "Bachelor of Science" when the user requires "Bachelor of
    Arts") gets rewritten before the cover letter / email reaches the user.

    Users without `exact_phrasing_required` notes get no patches — the prompt
    instructions are then the only enforcement layer.
    """
    patches = []
    for entry in profile.education:
        required = (entry.get("exact_phrasing_required") or "").strip()
        degree = (entry.get("degree") or "").strip()
        if not required or not degree:
            continue
        # If the required note explicitly mentions a wrong-degree variant to
        # avoid (free-form text), users can encode patterns like
        # "NEVER say 'Computer Engineering'". We extract any quoted phrases.
        wrong_variants = re.findall(r"['\"]([^'\"]+)['\"]", required)
        for variant in wrong_variants:
            try:
                patches.append((re.compile(re.escape(variant), re.IGNORECASE), degree))
            except re.error:
                continue
    return patches


_HARD_FACT_PATCHES = _build_hard_fact_patches(PROFILE)


def _enforce_hard_facts(text):
    """Rewrite any leaked factual mistakes into the correct phrasing.

    Runs after generation as a defense-in-depth guard in case a provider
    ignores the HARD FACTS instruction in the prompt. The list of patches
    is built from the user's profile (see _build_hard_fact_patches)."""
    if not text or not _HARD_FACT_PATCHES:
        return text
    for pattern, replacement in _HARD_FACT_PATCHES:
        text = pattern.sub(replacement, text)
    return text


def _strip_stray_markdown(text):
    """Remove markdown formatting tokens that would render literally in PDF/Word.

    The reference cover-letter PDFs the user wants to match use no bold,
    italic, or heading markup — labels in bullet sections are just plain
    "Label: explanation" prose. If the AI slips in "**Bold**", "__italic__",
    "# heading", or backticks, unwrap them here before any renderer sees
    them so no stray asterisks leak onto the final page.
    """
    if not text:
        return text
    # **bold** / __bold__  ->  bold   (keep inner text, drop the markers)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"__(.+?)__", r"\1", text, flags=re.DOTALL)
    # *italic* / _italic_  ->  italic   (only when surrounded by word chars
    # or line boundaries to avoid breaking e.g. "C*" code or bullet "- * x")
    text = re.sub(r"(?<![\*\w])\*([^\s\*][^\*]*?[^\s\*]|\S)\*(?!\w)", r"\1", text)
    text = re.sub(r"(?<![_\w])_([^\s_][^_]*?[^\s_]|\S)_(?!\w)", r"\1", text)
    # `code`  ->  code
    text = re.sub(r"`+([^`]+)`+", r"\1", text)
    # Leading "# heading" / "## heading"  ->  heading
    text = re.sub(r"(?m)^\s*#{1,6}\s+", "", text)
    return text


def _normalize_cover_letter_paragraphs(letter):
    """Defensive post-processor: if the AI returned a single huge paragraph,
    split it into proper paragraphs so the PDF/Word output looks like a real
    cover letter instead of one giant wall of text.

    Also strips stray markdown formatting (asterisks, underscores, backticks)
    so ``**Bold**`` never leaks into the final PDF.
    """
    text = _strip_stray_markdown((letter or "").strip())
    if not text:
        return text

    # If the AI already inserted blank-line paragraph breaks AND the result
    # has more than one block, keep its formatting as-is.
    blocks = re.split(r"\n\s*\n", text)
    if len(blocks) >= 3:
        return text

    # Otherwise, try to recover structure: collapse all whitespace, then
    # heuristically split into paragraphs at known boundaries.
    flat = re.sub(r"\s+", " ", text).strip()

    # 1) Split off the "Dear/To ... Team," greeting as its own block.
    greeting_match = re.match(r"^((?:Dear|To)\s+[^,]{2,80},)\s*", flat)
    greeting = ""
    if greeting_match:
        greeting = greeting_match.group(1)
        flat = flat[greeting_match.end():].strip()

    # 2) Pull the sign-off off the end if present. We look for the candidate's
    #    configured sign-off name (escaped for regex) so the recovery still
    #    works even when each user has a different name.
    signoff_name = _safe_signoff_name()
    name_pattern = re.escape(signoff_name) if signoff_name else None
    signoff = ""
    if name_pattern:
        signoff_match = re.search(
            rf"\b(Sincerely|Best regards|Warm regards|Thank you[^.]*)\.?\s*[,]?\s*{name_pattern}\.?\s*$",
            flat, flags=re.IGNORECASE,
        )
        if signoff_match:
            lead = signoff_match.group(1).strip().rstrip(",.")
            signoff = (
                f"{lead},\n{signoff_name}"
                if lead.lower() != signoff_name.lower()
                else signoff_name
            )
            flat = flat[:signoff_match.start()].strip()
        else:
            tail_name = re.search(rf"\b{name_pattern}\.?\s*$", flat, flags=re.IGNORECASE)
            if tail_name:
                signoff = signoff_name
                flat = flat[:tail_name.start()].strip()

    # 3) Split the bullet section if any inline "- " markers leaked into prose.
    bullet_section = ""
    bullet_header_match = re.search(
        r"(Why\s+I\s+am[^:]{0,120}:|Why\s+I\s+am\s+a\s+strong\s+fit[^:]{0,80}:)",
        flat, flags=re.IGNORECASE,
    )
    if bullet_header_match:
        header = bullet_header_match.group(1).strip()
        before = flat[:bullet_header_match.start()].strip()
        after = flat[bullet_header_match.end():].strip()
        # Pull out 2-4 bullet items from the "after" chunk.
        bullets = re.findall(r"-\s*([^-][^-\n]{20,300}?)(?=\s+-\s|$)", " " + after + " -")
        if bullets:
            bullet_lines = "\n".join(f"- {b.strip()}" for b in bullets[:4])
            bullet_section = f"{header}\n{bullet_lines}"
            # Whatever comes after the last bullet becomes the closing prose.
            tail_after_bullets = after.split("-")[-1].strip()
            flat = before
            if tail_after_bullets and tail_after_bullets not in bullets[-1]:
                flat = (flat + " " + tail_after_bullets).strip()

    # 4) Split remaining prose into ~3-4 paragraphs at sentence boundaries,
    #    targeting roughly 3 sentences per paragraph.
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", flat)
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        body_blocks = []
    else:
        target_paragraphs = 3 if len(sentences) <= 9 else 4
        per = max(2, math.ceil(len(sentences) / target_paragraphs))
        body_blocks = []
        for i in range(0, len(sentences), per):
            chunk = " ".join(sentences[i:i + per]).strip()
            if chunk:
                body_blocks.append(chunk)

    out_blocks = []
    if greeting:
        out_blocks.append(greeting)
    out_blocks.extend(body_blocks)
    if bullet_section:
        out_blocks.append(bullet_section)
    if signoff:
        out_blocks.append(signoff)

    return "\n\n".join(out_blocks).strip()


def _build_emergency_cover_letter(company, position, resume_code):
    company_text = (company or "your company").strip() or "your company"
    role_text = (position or "the role").strip() or "the role"
    signoff_name = _safe_signoff_name() or "Candidate"
    return (
        f"Dear {company_text} Hiring Team,\n\n"
        f"I am excited to apply for the {role_text} opportunity at {company_text}. "
        f"My background combines analytics, data engineering, and cross-functional execution, "
        f"which aligns well with the scope of this role. I have built data pipelines that handle "
        f"large daily volumes, shipped reporting work that influenced decisions, and collaborated "
        f"with stakeholders to translate business goals into measurable outcomes.\n\n"
        f"I am confident that the {resume_code} persona on my resume maps directly to the problems your "
        f"team is solving. I would welcome the chance to discuss how I can contribute to your next phase of growth.\n\n"
        f"Sincerely,\n{signoff_name}"
    )


@app.route("/api/cover-letter", methods=["POST"])
def draft_cover_letter():
    data = request.json or {}
    jd = data.get("jd", "")
    company = data.get("company", "the company")
    position = data.get("position", "the open position")
    resume_code = data.get("resume_code", DEFAULT_RESUME_CODE)
    tone = data.get("tone", "professional")

    if not jd.strip():
        return jsonify({"error": "Job Description missing"}), 400
    if resume_code not in RESUMES:
        resume_code = DEFAULT_RESUME_CODE

    prompt = _build_cover_letter_prompt(company, position, resume_code, jd, tone=tone)
    started = time.time()

    def _on_attempt_error(p, pr, err, lat_ms):
        _log_ai_usage("/api/cover-letter", p, "unknown", pr, "", lat_ms, status="error", error_code=str(err))

    try:
        response_text, model_used, provider, _, _ = _run_provider_chain(
            task="cover",
            chain=TASK_CHAIN_COVER,
            prompt_builder=lambda _p: prompt,
            expect_json=True,
            on_error_log=_on_attempt_error,
        )
        response_text = _strip_markdown_fences(response_text)
        try:
            result = json.loads(response_text)
            letter = (result.get("cover_letter") or "").strip()
        except Exception:
            # Some models return prose directly when ignoring JSON mode — keep it.
            letter = response_text.strip()
        if not letter:
            raise Exception("AI did not return a cover letter")
        # Defensive: if the model returned one giant blob, recover paragraph
        # structure so the PDF/Word output looks like a real letter. Also
        # strips stray markdown tokens like **bold** that would otherwise
        # render literally on the page.
        letter = _normalize_cover_letter_paragraphs(letter)
        # Fix any leaked factual errors (e.g. wrong degree title) using the
        # profile-driven hard-fact patches.
        letter = _enforce_hard_facts(letter)
        # Append the candidate's portfolio + LinkedIn URLs at the bottom
        # (idempotent — won't duplicate if the AI already included them).
        letter = _append_contact_urls_to_letter(letter)
        _log_ai_usage(
            "/api/cover-letter", provider, model_used, prompt, response_text,
            (time.time() - started) * 1000,
        )
        return jsonify({
            "cover_letter": letter,
            "company": company,
            "position": position,
            "resume_code": resume_code,
            "provider_used": provider,
        })
    except Exception as last_err:
        emergency = _build_emergency_cover_letter(company, position, resume_code)
        return jsonify({
            "cover_letter": emergency,
            "company": company,
            "position": position,
            "resume_code": resume_code,
            "provider_used": "template_fallback",
            "warning": f"AI providers unavailable: {last_err}",
        })


def _sanitize_filename_component(value, default="Company"):
    """Title-case + underscore: 'Acme Life Sciences' -> 'Acme_Life_Sciences'."""
    text = (value or "").strip()
    if not text:
        return default
    parts = re.findall(r"[A-Za-z0-9]+", text)
    if not parts:
        return default
    titled = [p[:1].upper() + p[1:].lower() if len(p) > 1 else p.upper() for p in parts]
    return "_".join(titled)


# ---------- Cover-letter PDF rendering (Noto Serif → Times-Roman fallback) ----
_FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
_NOTO_REG_PATH = os.path.join(_FONTS_DIR, "NotoSerif-Regular.ttf")
_NOTO_BOLD_PATH = os.path.join(_FONTS_DIR, "NotoSerif-Bold.ttf")
_PDF_FONTS_REGISTERED = False
_PDF_BODY_FONT = "Times-Roman"
_PDF_BOLD_FONT = "Times-Bold"


def _ensure_pdf_fonts():
    """Register Noto Serif if available; cache the choice across requests."""
    global _PDF_FONTS_REGISTERED, _PDF_BODY_FONT, _PDF_BOLD_FONT
    if _PDF_FONTS_REGISTERED:
        return
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        if os.path.exists(_NOTO_REG_PATH) and os.path.exists(_NOTO_BOLD_PATH):
            pdfmetrics.registerFont(TTFont("NotoSerif", _NOTO_REG_PATH))
            pdfmetrics.registerFont(TTFont("NotoSerif-Bold", _NOTO_BOLD_PATH))
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            registerFontFamily(
                "NotoSerif", normal="NotoSerif", bold="NotoSerif-Bold",
                italic="NotoSerif", boldItalic="NotoSerif-Bold",
            )
            _PDF_BODY_FONT = "NotoSerif"
            _PDF_BOLD_FONT = "NotoSerif-Bold"
    except Exception as e:
        # Fall back to built-in Times — visually 95% the same serif feel.
        print(f"PDF font registration fell back to Times-Roman: {e}")
    _PDF_FONTS_REGISTERED = True


def _render_cover_letter_pdf_bytes(letter_text):
    """Render a cover letter to A4 PDF bytes matching the reference style."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    _ensure_pdf_fonts()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
        title="Cover Letter",
        author=_safe_signoff_name() or _safe_candidate_full_name(),
    )

    body_style = ParagraphStyle(
        name="BodyText",
        fontName=_PDF_BODY_FONT,
        fontSize=11,
        leading=15,  # ~1.36 line height, matches LibreOffice default
        spaceAfter=10,
        alignment=TA_LEFT,
        firstLineIndent=0,
    )
    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=body_style,
        leftIndent=18,
        firstLineIndent=0,
        bulletIndent=4,
        spaceAfter=6,
    )

    story = []
    # Split into paragraphs on blank lines; preserve single-line breaks within a paragraph.
    blocks = [b for b in re.split(r"\n\s*\n", (letter_text or "").strip()) if b.strip()]

    bullet_prefix_re = re.compile(r"^\s*[\u2022\u25CF\u25AA\-\*]\s+")
    for block in blocks:
        lines = [ln.rstrip() for ln in block.split("\n") if ln.strip()]
        # If every line in this block starts with a bullet marker, render as a bullet list.
        if lines and all(bullet_prefix_re.match(ln) for ln in lines):
            for ln in lines:
                stripped = bullet_prefix_re.sub("", ln).strip()
                # Allow inline bold via **text** so the AI can emphasize key phrases naturally.
                stripped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
                story.append(Paragraph(stripped, bullet_style, bulletText="\u25CF"))
        else:
            text = "<br/>".join(lines)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(text, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


@app.route("/api/cover-letter/pdf", methods=["POST"])
def cover_letter_pdf():
    data = request.json or {}
    letter = (data.get("cover_letter") or "").strip()
    company = data.get("company", "").strip() or "Company"
    if not letter:
        return jsonify({"error": "Cover letter text missing"}), 400
    letter = _enforce_hard_facts(_strip_stray_markdown(letter))
    try:
        pdf_bytes = _render_cover_letter_pdf_bytes(letter)
    except Exception as e:
        return jsonify({"error": f"Failed to render PDF: {e}"}), 500

    filename = f"Cover_Letter_{_sanitize_filename_component(company)}.pdf"
    return send_file(
        io.BytesIO(pdf_bytes),
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf",
    )


@app.route("/api/cover-letter/docx", methods=["POST"])
def cover_letter_docx():
    data = request.json or {}
    letter = (data.get("cover_letter") or "").strip()
    company = data.get("company", "").strip() or "company"
    if not letter:
        return jsonify({"error": "Cover letter text missing"}), 400
    letter = _enforce_hard_facts(_strip_stray_markdown(letter))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Render paragraph-by-paragraph so double newlines form real paragraphs.
    for block in re.split(r"\n\s*\n", letter):
        cleaned = block.strip("\n").rstrip()
        if not cleaned:
            continue
        p = doc.add_paragraph()
        lines = cleaned.split("\n")
        for idx, line in enumerate(lines):
            if idx > 0:
                p.add_run().add_break()
            p.add_run(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Cover_Letter_{_sanitize_filename_component(company)}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/scout/jobs", methods=["POST"])
def scout_jobs():
    """
    Pull recent postings (LinkedIn jobs search + optional Greenhouse actor).

    Keyword resolution (first match wins):
      1) Request body "keyword" (manual override)
      2) If features.enable_scout_from_resumes: read all resume PDFs → AI → primary_linkedin_search
      3) SCOUT_JOB_KEYWORD_DEFAULT in .env

    LinkedIn time window uses f_TPR=r<seconds> (see SCOUT_LINKEDIN_POSTED_SECONDS).
    """
    data = request.json or {}
    manual_kw = (data.get("keyword") or "").strip()
    env_kw = (os.getenv("SCOUT_JOB_KEYWORD_DEFAULT") or "").strip()
    resume_driven_meta = None
    keyword = ""
    keyword_source = ""

    if manual_kw:
        keyword = manual_kw
        keyword_source = "manual"
    elif APP_SETTINGS.feature("enable_scout_from_resumes", False):
        blocks, extraction_errors = _gather_resume_keyword_blocks()
        if not blocks:
            return jsonify({
                "error": (
                    "Resume-driven scout needs readable PDF text from your personas. "
                    "Check resumes/ + config/resume_personas.json and install pypdf."
                ),
                "extraction_errors": extraction_errors,
            }), 400
        try:
            normalized, _rt, prov, mod = _run_resume_keyword_ai(
                blocks,
                usage_log_endpoint="/api/scout/jobs",
            )
        except Exception as err:
            msg, code = _friendly_ai_error(err, "AI could not derive search keywords from resumes")
            return jsonify({"error": msg, "extraction_errors": extraction_errors}), code
        keyword = _primary_linkedin_search_pick(normalized)
        if not keyword:
            return jsonify({
                "error": "AI response had no usable primary_linkedin_search or fallback phrases.",
                "extraction_errors": extraction_errors,
                "resume_ai_preview": {
                    "notes": normalized.get("notes"),
                    "combined_top_keywords": normalized.get("combined_top_keywords"),
                },
            }), 502
        resume_driven_meta = {
            "primary_linkedin_search": keyword,
            "personas_analyzed": [b["code"] for b in blocks],
            "extraction_errors": extraction_errors,
            "provider": prov,
            "model": mod,
            "ai_notes": normalized.get("notes"),
            "combined_top_keywords": normalized.get("combined_top_keywords"),
            "combined_role_families": normalized.get("combined_role_families"),
            "per_persona": normalized.get("per_persona"),
        }
        keyword_source = "resume"
    elif env_kw:
        keyword = env_kw
        keyword_source = "env_default"
    else:
        return jsonify(
            {
                "error": (
                    "No search keywords: type some in Step 1, set SCOUT_JOB_KEYWORD_DEFAULT in .env, "
                    "or enable features.enable_scout_from_resumes in config/app_settings.json "
                    "(uses your resume PDFs + AI)."
                )
            }
        ), 400
    try:
        max_items = max(1, min(int(data.get("max_items") or os.getenv("SCOUT_JOB_MAX_ITEMS", "80")), 10000))
    except Exception:
        max_items = 80
    try:
        seconds_filter = max(60, min(int(os.getenv("SCOUT_LINKEDIN_POSTED_SECONDS", "10000")), 31536000))
    except Exception:
        seconds_filter = 10000

    linkedin_actor = (
        (os.getenv("APIFY_ACTOR_LINKEDIN_JOBS_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_JOB_SCOUT_LINKEDIN_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_DISCOVERY_ID") or "").strip()
    )
    greenhouse_actor = (
        (os.getenv("APIFY_ACTOR_GREENHOUSE_JOBS_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_GREENHOUSE_ID") or "").strip()
    )

    if not linkedin_actor:
        return jsonify({"error": "Configure APIFY_ACTOR_LINKEDIN_JOBS_ID (or reuse APIFY_ACTOR_DISCOVERY_ID)."}), 503

    scout_log_path = Path(PROJECT_ROOT) / "data" / "job_scout_runs.jsonl"

    warnings: list = []

    linkedin_url = _build_linkedin_jobs_search_url(keyword, seconds_filter)

    extras_li = _parse_optional_json_env("APIFY_LINKEDIN_JOBS_INPUT_JSON")
    linkedin_input = _linkedin_jobs_actor_input(
        linkedin_url=linkedin_url,
        keyword=keyword,
        max_items=max_items,
        actor_id=linkedin_actor,
        extras=extras_li,
    )

    extras_gh = _parse_optional_json_env("APIFY_GREENHOUSE_JOBS_INPUT_JSON")

    combined: list = []
    results_meta: dict = {}

    try:
        service = ApifyClientService()
    except ApifyConfigError as cfg_err:
        return jsonify({"error": str(cfg_err)}), 503

    try:
        li_res = service.run_actor(
            actor_id=linkedin_actor,
            actor_input=linkedin_input,
            log_path=scout_log_path,
            record_type="job_scout_linkedin",
            extra_meta={"keyword": keyword, "linkedin_search_url": linkedin_url},
        )
        results_meta["linkedin"] = {
            "run_id": li_res.run_id,
            "item_count": len(li_res.items or []),
            "estimated_cost_usd": li_res.estimated_cost_usd,
        }
        combined.extend(_extract_job_candidates_from_items("linkedin", li_res.items or []))
    except ApifySafeStop as safe_stop:
        return jsonify(
            {
                "error": str(safe_stop),
                "error_type": "budget_safe_stop",
                "paused": True,
            }
        ), 429
    except Exception as err:
        msg, code = _normalize_apify_error(err)
        warnings.append(f"LinkedIn scout failed: {msg}")
        results_meta["linkedin"] = {"error": msg}

    if greenhouse_actor:
        if extras_gh:
            gh_input = _deep_merge_dicts(
                extras_gh,
                {
                    "maxItems": max_items,
                    "keyword": keyword,
                    "keywords": keyword,
                    "search": keyword,
                    "query": keyword,
                },
            )
        else:
            gh_input = {
                "keyword": keyword,
                "keywords": keyword,
                "search": keyword,
                "query": keyword,
                "maxItems": max_items,
            }
        try:
            gh_res = service.run_actor(
                actor_id=greenhouse_actor,
                actor_input=gh_input,
                log_path=scout_log_path,
                record_type="job_scout_greenhouse",
                extra_meta={"keyword": keyword},
            )
            results_meta["greenhouse"] = {
                "run_id": gh_res.run_id,
                "item_count": len(gh_res.items or []),
                "estimated_cost_usd": gh_res.estimated_cost_usd,
            }
            combined.extend(_extract_job_candidates_from_items("greenhouse", gh_res.items or []))
        except ApifySafeStop as safe_stop:
            return jsonify(
                {
                    "error": str(safe_stop),
                    "error_type": "budget_safe_stop",
                    "paused": True,
                    "partial_jobs": _rank_scout_jobs_by_personas(_merge_jobs_by_url(combined)),
                    "warnings": warnings,
                }
            ), 429
        except Exception as err:
            msg, _code = _normalize_apify_error(err)
            warnings.append(f"Greenhouse scout failed: {msg}")
            results_meta["greenhouse"] = {"error": msg}
    else:
        warnings.append("Greenhouse scout skipped — set APIFY_ACTOR_GREENHOUSE_JOBS_ID in .env.")

    merged = _rank_scout_jobs_by_personas(_merge_jobs_by_url(combined))

    payload = {
        "success": True,
        "keyword": keyword,
        "keyword_source": keyword_source,
        "linkedin_search_url": linkedin_url,
        "posted_within_seconds": seconds_filter,
        "max_items_per_source": max_items,
        "jobs": merged,
        "runs": results_meta,
        "warnings": warnings,
        "ranking_note": (
            "Jobs are ordered by overlap with your resume_personas triggers/stacks (local text match, "
            "no extra AI calls). Hover the % chip for why; Deep Strategic Analysis is the real fit score."
        ),
    }
    if resume_driven_meta is not None:
        payload["resume_driven_meta"] = resume_driven_meta
    return jsonify(payload)


@app.route("/api/experiment/resume-keywords", methods=["POST"])
def experiment_resume_keywords():
    """
    EXPERIMENT: read all configured resume PDFs + persona metadata, call the AI
    to suggest LinkedIn search phrases and target roles. Disable via
    config/app_settings.json -> features.enable_resume_keyword_experiment.
    """
    if not APP_SETTINGS.feature("enable_resume_keyword_experiment", False):
        return jsonify({
            "error": (
                "Resume keyword experiment is disabled. Set "
                "features.enable_resume_keyword_experiment to true in "
                "config/app_settings.json and restart the app."
            ),
        }), 404

    blocks, extraction_errors = _gather_resume_keyword_blocks()
    if not blocks:
        return jsonify({
            "error": (
                "No resume text could be extracted. Ensure PDFs exist under resumes/ "
                "and paths in config/resume_personas.json match. Install pypdf: pip install pypdf"
            ),
            "extraction_errors": extraction_errors,
        }), 400

    try:
        result, _response_dump, provider, model_used = _run_resume_keyword_ai(
            blocks,
            usage_log_endpoint="/api/experiment/resume-keywords",
        )
        result["success"] = True
        result["extraction_errors"] = extraction_errors
        result["personas_analyzed"] = [b["code"] for b in blocks]
        result["_experiment"] = True
        result["provider"] = provider
        result["model"] = model_used
        return jsonify(result)
    except Exception as err:
        msg, code = _friendly_ai_error(err, "Resume keyword suggestions failed")
        return jsonify({"error": msg, "extraction_errors": extraction_errors}), code


@app.route("/api/discovery", methods=["POST"])
def discovery():
    data = request.json or {}
    job_id = (data.get("job_id") or "").strip()
    source_url = (data.get("source_url") or "").strip()
    if not job_id:
        return jsonify({"error": "job_id is required for discovery deduplication."}), 400
    if not source_url:
        return jsonify({"error": "source_url is required for discovery."}), 400

    try:
        service = ApifyClientService()
        result = service.run_discovery(
            job_id=job_id,
            source_url=source_url,
            actor_input=_build_discovery_input(data),
        )
        return jsonify({
            "success": True,
            "status": result.status,
            "skipped": result.skipped,
            "reason": result.reason,
            "run_id": result.run_id,
            "dataset_id": result.default_dataset_id,
            "item_count": len(result.items or []),
            "items": result.items,
            "estimated_cost_usd": result.estimated_cost_usd,
        })
    except ApifyConfigError as err:
        return jsonify({"error": str(err), "error_type": "apify_config"}), 503
    except ApifySafeStop as safe_stop:
        return jsonify({
            "error": str(safe_stop),
            "error_type": "budget_safe_stop",
            "paused": True,
        }), 429
    except Exception as err:
        msg, code = _normalize_apify_error(err)
        return jsonify({"error": msg}), code


@app.route("/api/enrichment", methods=["POST"])
def enrichment():
    data = request.json or {}
    actor_input = data.get("actor_input")
    if not isinstance(actor_input, dict):
        actor_input = {
            "company": (data.get("company") or "").strip(),
            "position": (data.get("position") or "").strip(),
            "contact_name": (data.get("contact_name") or "").strip(),
            "source_url": (data.get("source_url") or "").strip(),
            "discovery_run_id": (data.get("discovery_run_id") or "").strip(),
        }

    try:
        service = ApifyClientService()
        result = service.run_enrichment(actor_input=actor_input)
        handoff = _extract_enrichment_handoff(result.items or [])
        return jsonify({
            "success": True,
            "status": result.status,
            "run_id": result.run_id,
            "dataset_id": result.default_dataset_id,
            "item_count": len(result.items or []),
            "items": result.items,
            "handoff": handoff,
            "estimated_cost_usd": result.estimated_cost_usd,
        })
    except ApifyConfigError as err:
        return jsonify({"error": str(err), "error_type": "apify_config"}), 503
    except ApifySafeStop as safe_stop:
        return jsonify({
            "error": str(safe_stop),
            "error_type": "budget_safe_stop",
            "paused": True,
        }), 429
    except Exception as err:
        msg, code = _normalize_apify_error(err)
        return jsonify({"error": msg}), code


@app.route("/api/send", methods=["POST"])
def send():
    _ensure_scheduler_started()
    data = request.json
    target_email = data.get("email")
    subject = data.get("subject")
    body = data.get("body")
    resume_code = data.get("resume_code")
    company = data.get("company")
    position = data.get("position")
    contact_name = data.get("contact_name", "N/A")
    schedule_next_day_8am = bool(data.get("schedule_next_day_8am"))
    
    if not target_email or not subject or not body or not resume_code or not company:
        return jsonify({"error": "Missing required fields"}), 400
        
    selected_pdf = RESUMES.get(resume_code)
    if not selected_pdf:
        return jsonify({"error": "Invalid resume code"}), 400
        
    sender_ok = bool((os.getenv("SENDER_EMAIL") or "").strip() and (os.getenv("EMAIL_APP_PASSWORD") or "").strip())

    try:
        if schedule_next_day_8am:
            if not ENABLE_EMAIL_SCHEDULING:
                return jsonify({"error": "Email scheduling is disabled in app_settings.json."}), 400
            if not sender_ok:
                return jsonify({
                    "error": (
                        "Cannot schedule email: set SENDER_EMAIL and EMAIL_APP_PASSWORD in .env "
                        "(Gmail App Password). Scheduled sends use the same SMTP credentials."
                    ),
                }), 400
            send_at = _next_day_send_time_local()
            _queue_scheduled_email({
                "id": str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(timespec="seconds"),
                "send_at": send_at.isoformat(timespec="seconds"),
                "status": "pending",
                "email": target_email,
                "subject": subject,
                "body": body,
                "attachment_path": selected_pdf,
                "resume_code": resume_code,
                "company": company,
                "position": position,
                "contact_name": contact_name,
            })
            log_job(
                company=company,
                position=position,
                resume_used=resume_code,
                contact_name=contact_name,
                role="N/A",
                platform="Web Scheduled Email",
                message=subject,
                status=f"Scheduled for {send_at.strftime('%Y-%m-%d %H:%M')}",
                tracker_path=JOB_TRACKER_PATH,
            )
            return jsonify({
                "success": True,
                "scheduled": True,
                "send_at": send_at.isoformat(timespec="seconds"),
            })
        if not ENABLE_EMAIL_SENDING:
            return jsonify({"error": "Email sending is disabled in app_settings.json. Set features.enable_email_sending=true to enable."}), 400
        if not sender_ok:
            return jsonify({
                "error": (
                    "Gmail SMTP not configured: set SENDER_EMAIL to your Gmail address and "
                    "EMAIL_APP_PASSWORD to a Google App Password (not your normal login). "
                    "See .env.example — then restart the app."
                ),
            }), 400
        send_email(to_email=target_email, subject=subject, body=body, attachment_path=selected_pdf)
        log_job(
            company=company,
            position=position,
            resume_used=resume_code,
            contact_name=contact_name,
            role="N/A",
            platform="Web Direct Email",
            message=subject,
            status="Applied",
            tracker_path=JOB_TRACKER_PATH,
        )
        return jsonify({"success": True, "scheduled": False})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/config", methods=["GET"])
def app_config():
    """Lightweight config blob exposing UI-relevant settings + profile metadata.
    The frontend reads this on load to wire up affinity buttons, signoff names,
    feature flags, etc., without ever embedding personal info in HTML/JS files."""
    actor_li = (
        (os.getenv("APIFY_ACTOR_LINKEDIN_JOBS_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_JOB_SCOUT_LINKEDIN_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_DISCOVERY_ID") or "").strip()
    )
    greenhouse_id = (
        (os.getenv("APIFY_ACTOR_GREENHOUSE_JOBS_ID") or "").strip()
        or (os.getenv("APIFY_ACTOR_GREENHOUSE_ID") or "").strip()
    )
    apify_token = (os.getenv("APIFY_TOKEN") or "").strip()
    discovery_actor_only = (os.getenv("APIFY_ACTOR_DISCOVERY_ID") or "").strip()
    enrichment_actor_only = (os.getenv("APIFY_ACTOR_ENRICHMENT_ID") or "").strip()
    return jsonify({
        "ui": {
            "app_title": APP_SETTINGS.ui_text("app_title", "Career Command Center"),
            "app_subtitle": APP_SETTINGS.ui_text("app_subtitle", "AI-Powered Job Fit & Outreach Engine"),
            "schedule_hour_label": f"{EMAIL_SCHEDULER_HOUR_LOCAL:02d}:00",
        },
        "candidate": {
            "full_name": _safe_candidate_full_name(),
            "preferred_name": PROFILE.preferred_name,
            "signoff_name": _safe_signoff_name(),
        },
        "features": {
            "enable_followup_agent": ENABLE_FOLLOWUP_AGENT,
            "enable_email_sending": ENABLE_EMAIL_SENDING,
            "enable_email_scheduling": ENABLE_EMAIL_SCHEDULING,
            "enable_alumni_search_button": (
                PROFILE.alumni_search_enabled
                and APP_SETTINGS.feature("enable_alumni_search_button", True)
            ),
            "enable_resume_keyword_experiment": APP_SETTINGS.feature(
                "enable_resume_keyword_experiment",
                False,
            ),
            "enable_scout_from_resumes": APP_SETTINGS.feature(
                "enable_scout_from_resumes",
                False,
            ),
        },
        "alumni": {
            "enabled": PROFILE.alumni_search_enabled,
            "label": PROFILE.alumni_button_label,
            "emoji": PROFILE.alumni_button_emoji,
            "school_slug": PROFILE.alumni_school_slug,
        },
        # Job scout: documented defaults so clones know which Actors this template targets.
        "job_scout": {
            "default_keyword_hint": (os.getenv("SCOUT_JOB_KEYWORD_DEFAULT") or "").strip(),
            "linkedin_actor_recommended_id": "curious_coder/linkedin-jobs-scraper",
            "linkedin_actor_store_url": "https://apify.com/curious_coder/linkedin-jobs-scraper",
            "linkedin_actor_configured_id": actor_li or None,
            "greenhouse_actor_configured_id": greenhouse_id or None,
        },
        # Step 3 optional auto-discovery + enrichment (separate from Step 1 job scout actors).
        "apify_step3": {
            "token_configured": bool(apify_token),
            "discovery_actor_configured": bool(discovery_actor_only),
            "enrichment_actor_configured": bool(enrichment_actor_only),
            "discovery_enrichment_ready": bool(
                apify_token and discovery_actor_only and enrichment_actor_only
            ),
        },
    })


@app.route("/api/resumes", methods=["GET"])
def list_resumes():
    """Expose the registered resume personas to the frontend so both the
    Step-3 'Optimum persona' dropdown and the Step-4 'Attach' toggle stay
    in sync with whatever is in RESUMES/RESUME_LABELS. Also reports whether
    the underlying PDF actually exists on disk — so a stale entry (pointing
    at a missing file) can be surfaced in the UI instead of failing silently
    at send-time."""
    items = []
    for code, filename in RESUMES.items():
        items.append({
            "code": code,
            "label": RESUME_LABELS.get(code, code),
            "filename": filename,
            "exists": os.path.exists(filename),
        })
    return jsonify({"resumes": items})


@app.route("/api/quota-health", methods=["GET"])
def quota_health():
    stats = _read_usage_summary_today()
    analyze_today = stats["analyze_today"]
    near_soft = analyze_today >= max(1, int(DAILY_ANALYZE_SOFT_CAP * 0.8))
    over_soft = analyze_today >= DAILY_ANALYZE_SOFT_CAP
    over_hard = analyze_today >= DAILY_ANALYZE_HARD_CAP
    return jsonify({
        "analyze_today": analyze_today,
        "soft_cap": DAILY_ANALYZE_SOFT_CAP,
        "hard_cap": DAILY_ANALYZE_HARD_CAP,
        "near_soft": near_soft,
        "over_soft": over_soft,
        "over_hard": over_hard,
        "quota_429_today": stats["quota_429_today"],
        "requests_today": stats["requests_today"],
        "gemini_today": stats["gemini_today"],
        "ollama_today": stats["ollama_today"],
        "groq_today": stats.get("groq_today", 0),
        "cerebras_today": stats.get("cerebras_today", 0),
        "github_today": stats.get("github_today", 0),
        "providers_configured": {
            "groq": bool(GROQ_API_KEY),
            "cerebras": bool(CEREBRAS_API_KEY),
            "gemini": bool(api_key),
            "github": bool(GITHUB_MODELS_TOKEN),
            "ollama": True,
        },
        "chains": {
            "analyze": TASK_CHAIN_ANALYZE,
            "draft": TASK_CHAIN_DRAFT,
            "note": TASK_CHAIN_NOTE,
            "ask": TASK_CHAIN_ASK,
            "cover": TASK_CHAIN_COVER,
        },
    })

if __name__ == "__main__":
    # macOS Monterey+ reserves port 5000 for AirPlay Receiver, which returns HTTP 403
    # and intermittently hijacks localhost requests. Default Flask to 5001; override via FLASK_PORT.
    flask_port = int(os.getenv("FLASK_PORT", "5001"))
    flask_host = os.getenv("FLASK_HOST", "127.0.0.1")
    app.run(debug=True, host=flask_host, port=flask_port)
